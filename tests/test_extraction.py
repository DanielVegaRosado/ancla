"""Tests for deterministic text extraction (PDF and .docx).

Uses real files generated within the test itself — a byte-for-byte valid
PDF, a real .docx via python-docx — instead of mocking the libraries: what
matters here is that extraction is *faithful*, and only going through the
real parser proves that.
"""
from __future__ import annotations

import io
import sys

import pytest

from ancla.profile.extraction import ExtractionError, extract_text


def _pdf_minimo(texto: str) -> bytes:
    """A valid one-page PDF with `texto` as extractable content, built with
    exact (not approximate) xref offsets."""
    contenido = f"BT /F1 12 Tf 10 50 Td ({texto}) Tj ET".encode("latin-1")
    objetos = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>"
        b"/MediaBox[0 0 300 100]/Contents 5 0 R>>",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
        b"<</Length " + str(len(contenido)).encode() + b">>\nstream\n"
        + contenido + b"\nendstream",
    ]
    partes = [b"%PDF-1.4\n"]
    offsets = [0]
    for n, cuerpo in enumerate(objetos, start=1):
        offsets.append(sum(len(p) for p in partes))
        partes.append(f"{n} 0 obj".encode() + cuerpo + b"endobj\n")
    inicio_xref = sum(len(p) for p in partes)
    xref = [b"xref\n", f"0 {len(objetos) + 1}\n".encode(), b"0000000000 65535 f \n"]
    for offset in offsets[1:]:
        xref.append(f"{offset:010d} 00000 n \n".encode())
    partes += xref
    partes.append(
        f"trailer<</Size {len(objetos) + 1}/Root 1 0 R>>\n"
        f"startxref\n{inicio_xref}\n%%EOF".encode()
    )
    return b"".join(partes)


def _docx_con(parrafos: list[str]) -> bytes:
    import docx

    documento = docx.Document()
    for parrafo in parrafos:
        documento.add_paragraph(parrafo)
    buffer = io.BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def test_extrae_texto_de_un_pdf_real():
    datos = _pdf_minimo("Experiencia en Python y SQL, tres anhos en el sector.")
    texto = extract_text("cv.pdf", datos)
    assert "Experiencia en Python y SQL" in texto


def test_extrae_texto_de_un_docx_real():
    datos = _docx_con(
        [
            "Daniel Vega Rosado",
            "ML Developer — Telco Churn (2026 - actualidad)",
            "Pipeline completo de datos con Python y Optuna.",
        ]
    )
    texto = extract_text("cv.docx", datos)
    assert "ML Developer" in texto
    assert "Python y Optuna" in texto


def test_extrae_texto_de_tablas_en_docx():
    """A CV often puts the stack or dates in a table, not only in paragraphs."""
    import docx

    documento = docx.Document()
    documento.add_paragraph("Daniel Vega Rosado — Ingeniero de datos junior")
    tabla = documento.add_table(rows=1, cols=2)
    tabla.rows[0].cells[0].text = "Lenguajes: Python, SQL"
    tabla.rows[0].cells[1].text = "Herramientas: Optuna, Airflow"
    buffer = io.BytesIO()
    documento.save(buffer)

    texto = extract_text("cv.docx", buffer.getvalue())
    assert "Python, SQL" in texto and "Optuna, Airflow" in texto


def test_un_formato_no_soportado_lanza_error_claro():
    with pytest.raises(ExtractionError, match="cv.txt"):
        extract_text("cv.txt", b"cualquier cosa")


def test_sin_pypdf_instalado_lo_dice_en_vez_de_culpar_al_pdf(monkeypatch):
    """The real case that prompted this: Daniel had `pypdf` installed in a
    different environment from the one running `run.py`. `ModuleNotFoundError`
    is an `Exception` like any other, so the generic safety net caught it —
    but with the "your PDF might be corrupted" message, which is actively
    misleading when the real problem is a missing dependency. It has to
    show the same actionable message `ai/groq.py` already uses for the same
    case with the Groq SDK."""
    monkeypatch.setitem(sys.modules, "pypdf", None)

    with pytest.raises(ExtractionError, match="pip install -r requirements.txt"):
        extract_text("cv.pdf", b"lo que sea")


def test_sin_python_docx_instalado_lo_dice_tambien(monkeypatch):
    monkeypatch.setitem(sys.modules, "docx", None)

    with pytest.raises(ExtractionError, match="pip install -r requirements.txt"):
        extract_text("cv.docx", b"lo que sea")


def test_una_excepcion_no_prevista_del_extractor_no_llega_a_500(monkeypatch):
    """The real case that prompted this: a real PDF (protected, with unusual
    fonts, exported by a specific tool) can make pypdf fail with an
    exception other than `PdfReadError` — the only one caught inside
    `_text_from_pdf`. Since the view only catches `ExtractionError`, anything
    else used to bubble up uncontrolled all the way to Flask and come out as
    a 500. Here that "anything else" is simulated with a generic
    `RuntimeError`, so as not to depend on one specific, fragile pathological
    PDF."""
    import pypdf

    class BrokenReader:
        def __init__(self, *a, **kw):
            raise RuntimeError("fallo interno de pypdf no relacionado con el formato")

    monkeypatch.setattr(pypdf, "PdfReader", BrokenReader)

    with pytest.raises(ExtractionError, match="No se ha podido leer"):
        extract_text("cv.pdf", b"lo que sea, no llega a importar")


def test_un_pdf_invalido_lanza_error_en_vez_de_reventar():
    with pytest.raises(ExtractionError):
        extract_text("cv.pdf", b"esto no es un PDF de verdad")


def test_un_docx_invalido_lanza_error_en_vez_de_reventar():
    with pytest.raises(ExtractionError):
        extract_text("cv.docx", b"esto no es un docx de verdad")


def test_un_pdf_sin_texto_real_avisa_de_que_puede_ser_una_imagen():
    """Un PDF escaneado (imagen sin capa de texto) extrae vacío o casi vacío:
    hay que decirlo, no fingir que se leyó un CV en blanco."""
    datos = _pdf_minimo("")
    with pytest.raises(ExtractionError, match="escaneado"):
        extract_text("cv.pdf", datos)


def test_la_extension_no_distingue_mayusculas():
    datos = _pdf_minimo("Contenido de prueba suficientemente largo para pasar el minimo.")
    assert "Contenido de prueba" in extract_text("CV.PDF", datos)
