"""Tests for the Import screen: the import's ephemeral state and the HTTP
routes end to end (upload, review, save the selected ones, discard).
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest

from ancla.profile import store
from ancla.profile.model import Bilingual, Experience, SpokenLanguage, Skill
from ancla.web import create_app
from ancla.web import import_batch as modulo_importacion


@pytest.fixture
def cliente_web(tmp_path: Path):
    app = create_app(raiz_perfil=tmp_path / "perfil", settings_path=tmp_path / "ajustes.json")
    app.config["TESTING"] = True
    return app.test_client()


def _experiencia(id: str = "ml-dev") -> Experience:
    return Experience(
        id=id,
        title=Bilingual(es="ML Developer", en="ML Developer"),
        period=Bilingual(es="2026", en="2026"),
        bullets=Bilingual(es=["Pipeline completo"], en=["Full pipeline"]),
        stack=Bilingual(es="Python", en="Python"),
        keywords=["ml"],
    )


def _skill(id: str = "python") -> Skill:
    return Skill(id=id, name=Bilingual(es="Python", en="Python"), category="lenguaje", keywords=["py"])


def _skill_personal(id: str = "equipo") -> Skill:
    return Skill(id=id, name=Bilingual(es="Trabajo en equipo", en="Teamwork"), keywords=["team player"])


def _idioma(id: str = "ingles") -> SpokenLanguage:
    return SpokenLanguage(
        id=id,
        name=Bilingual(es="Inglés", en="English"),
        level=Bilingual(es="C1", en="C1"),
        keywords=["advanced english"],
    )


# --------------------------------------------------------------------------
# web/importacion.py: ephemeral state, round trip
# --------------------------------------------------------------------------


def test_sin_importacion_guardada_no_hay_nada_que_cargar(tmp_path: Path):
    assert modulo_importacion.load_import(tmp_path) is None


def test_guardar_y_cargar_importacion_hace_ida_y_vuelta(tmp_path: Path):
    original = modulo_importacion.ImportBatch(
        experiencias=[_experiencia()], skills=[_skill()], avisos=["ojo con esto"]
    )
    modulo_importacion.save_import(tmp_path, original)
    recargada = modulo_importacion.load_import(tmp_path)
    assert recargada == original


def test_borrar_importacion_la_deja_indisponible(tmp_path: Path):
    modulo_importacion.save_import(
        tmp_path, modulo_importacion.ImportBatch(experiencias=[_experiencia()])
    )
    modulo_importacion.delete_import(tmp_path)
    assert modulo_importacion.load_import(tmp_path) is None


def test_borrar_importacion_sin_fichero_no_falla(tmp_path: Path):
    modulo_importacion.delete_import(tmp_path)  # no debe lanzar


def test_un_fichero_de_importacion_corrupto_no_revienta(tmp_path: Path):
    (tmp_path / modulo_importacion.NOMBRE_FICHERO).write_text("esto no es JSON", encoding="utf-8")
    assert modulo_importacion.load_import(tmp_path) is None


# --------------------------------------------------------------------------
# HTTP routes
# --------------------------------------------------------------------------


def test_ver_importar(cliente_web):
    respuesta = cliente_web.get("/perfil/importar")
    assert respuesta.status_code == 200
    assert "Importar".encode("utf-8") in respuesta.data


def test_importar_sin_fichero_ni_texto_pide_uno_de_los_dos(cliente_web):
    respuesta = cliente_web.post("/perfil/importar", data={"texto": ""})
    assert respuesta.status_code == 200
    assert "Sube un fichero o pega el texto".encode("utf-8") in respuesta.data


def test_importar_sin_clave_configurada_redirige_a_ajustes(cliente_web):
    respuesta = cliente_web.post(
        "/perfil/importar", data={"texto": "Un CV cualquiera con suficiente texto."}
    )
    assert respuesta.status_code == 302
    assert respuesta.location.endswith("/ajustes")


def test_revisar_sin_importacion_pendiente_redirige_a_importar(cliente_web):
    respuesta = cliente_web.get("/perfil/importar/revisar", follow_redirects=True)
    assert "No hay ninguna importación pendiente".encode("utf-8") in respuesta.data


def test_revisar_muestra_las_candidatas_guardadas(cliente_web, tmp_path: Path):
    modulo_importacion.save_import(
        tmp_path / "perfil",
        modulo_importacion.ImportBatch(experiencias=[_experiencia()], skills=[_skill()]),
    )
    respuesta = cliente_web.get("/perfil/importar/revisar")
    assert "ML Developer".encode("utf-8") in respuesta.data
    assert "Python".encode("utf-8") in respuesta.data


def test_guardar_solo_lo_marcado(cliente_web, tmp_path: Path):
    root = tmp_path / "perfil"
    modulo_importacion.save_import(
        root,
        modulo_importacion.ImportBatch(
            experiencias=[_experiencia("uno"), _experiencia("dos")], skills=[_skill()]
        ),
    )
    cliente_web.post(
        "/perfil/importar/guardar",
        data={
            "exp-0": "1",
            "exp-0-titulo_es": "ML Developer",
            "exp-0-titulo_en": "ML Developer",
            "exp-0-periodo_es": "2026",
            "exp-0-periodo_en": "2026",
            "exp-0-bullets_es": "Pipeline completo",
            "exp-0-bullets_en": "Full pipeline",
            "exp-0-stack_es": "Python",
            "exp-0-stack_en": "Python",
            # "exp-1" no viene en el formulario: no estaba marcado
        },
    )
    perfil = store.load_profile(root)
    assert perfil.experience("uno") is not None
    assert perfil.experience("dos") is None
    assert perfil.skill("python") is None  # tampoco estaba marcado


def test_guardar_permite_editar_antes_de_confirmar(cliente_web, tmp_path: Path):
    root = tmp_path / "perfil"
    modulo_importacion.save_import(
        root, modulo_importacion.ImportBatch(skills=[_skill()])
    )
    cliente_web.post(
        "/perfil/importar/guardar",
        data={
            "skill-0": "1",
            "skill-0-nombre_es": "Python avanzado",
            "skill-0-nombre_en": "Advanced Python",
            "skill-0-categoria": "lenguaje",
        },
    )
    perfil = store.load_profile(root)
    assert perfil.skill("python").name["es"] == "Python avanzado"


def test_una_candidata_incompleta_no_se_guarda_pero_avisa(cliente_web, tmp_path: Path):
    """If the English name gets erased while editing, normal validation
    rejects it — just like in the manual form — instead of saving it broken."""
    root = tmp_path / "perfil"
    modulo_importacion.save_import(
        root, modulo_importacion.ImportBatch(skills=[_skill()])
    )
    respuesta = cliente_web.post(
        "/perfil/importar/guardar",
        data={"skill-0": "1", "skill-0-nombre_es": "Python", "skill-0-nombre_en": ""},
        follow_redirects=True,
    )
    perfil = store.load_profile(root)
    assert perfil.skill("python") is None
    assert "no se pudieron guardar".encode("utf-8") in respuesta.data


def test_guardar_una_skill_personal_importada_va_al_catalogo_correcto(cliente_web, tmp_path: Path):
    """Real case that prompted this: a "PERSONAL" section of the CV (Trabajo
    en equipo, Team player...) has to end up in personal_skills, never mixed
    in with technical skills."""
    root = tmp_path / "perfil"
    modulo_importacion.save_import(
        root, modulo_importacion.ImportBatch(skills_personales=[_skill_personal()])
    )
    cliente_web.post(
        "/perfil/importar/guardar",
        data={
            "skillpersonal-0": "1",
            "skillpersonal-0-nombre_es": "Trabajo en equipo",
            "skillpersonal-0-nombre_en": "Teamwork",
        },
    )
    perfil = store.load_profile(root)
    assert perfil.personal_skill("equipo") is not None
    assert perfil.skill("equipo") is None  # nunca en el catálogo técnico


def test_guardar_un_idioma_importado(cliente_web, tmp_path: Path):
    root = tmp_path / "perfil"
    modulo_importacion.save_import(
        root, modulo_importacion.ImportBatch(idiomas=[_idioma()])
    )
    cliente_web.post(
        "/perfil/importar/guardar",
        data={
            "idioma-0": "1",
            "idioma-0-nombre_es": "Inglés",
            "idioma-0-nombre_en": "English",
            "idioma-0-nivel_es": "C1 Avanzado",
            "idioma-0-nivel_en": "C1 Advanced",
        },
    )
    perfil = store.load_profile(root)
    assert perfil.language("ingles") is not None
    assert perfil.language("ingles").level["es"] == "C1 Avanzado"


def test_una_skill_personal_o_idioma_no_marcados_no_se_guardan(cliente_web, tmp_path: Path):
    root = tmp_path / "perfil"
    modulo_importacion.save_import(
        root,
        modulo_importacion.ImportBatch(
            skills_personales=[_skill_personal()], idiomas=[_idioma()]
        ),
    )
    cliente_web.post("/perfil/importar/guardar", data={})
    perfil = store.load_profile(root)
    assert perfil.personal_skill("equipo") is None
    assert perfil.language("ingles") is None


def test_revisar_muestra_skills_personales_e_idiomas(cliente_web, tmp_path: Path):
    modulo_importacion.save_import(
        tmp_path / "perfil",
        modulo_importacion.ImportBatch(
            skills_personales=[_skill_personal()], idiomas=[_idioma()]
        ),
    )
    respuesta = cliente_web.get("/perfil/importar/revisar")
    html = respuesta.data.decode("utf-8")
    assert "Trabajo en equipo" in html
    assert "Inglés" in html and "C1" in html


def test_guardar_limpia_la_importacion_pendiente(cliente_web, tmp_path: Path):
    root = tmp_path / "perfil"
    modulo_importacion.save_import(root, modulo_importacion.ImportBatch(skills=[_skill()]))
    cliente_web.post("/perfil/importar/guardar", data={"skill-0": "1", "skill-0-nombre_es": "Python", "skill-0-nombre_en": "Python"})
    assert modulo_importacion.load_import(root) is None


def test_guardar_sin_importacion_pendiente_redirige(cliente_web):
    respuesta = cliente_web.post("/perfil/importar/guardar", data={})
    assert respuesta.status_code == 302


def test_descartar_borra_la_importacion_sin_guardar_nada(cliente_web, tmp_path: Path):
    root = tmp_path / "perfil"
    modulo_importacion.save_import(root, modulo_importacion.ImportBatch(skills=[_skill()]))
    cliente_web.post("/perfil/importar/descartar")
    assert modulo_importacion.load_import(root) is None
    assert store.load_profile(root).skill("python") is None


def test_un_perfil_vacio_enlaza_a_importar_desde_mi_perfil(cliente_web):
    respuesta = cliente_web.get("/perfil")
    assert b'href="/perfil/importar"' in respuesta.data
    assert "¿Ya tienes un CV? Impórtalo".encode("utf-8") in respuesta.data


def test_un_perfil_con_datos_tambien_enlaza_a_importar_desde_mi_perfil(cliente_web, tmp_path: Path):
    """The link used to only appear with an empty profile; now it stays
    there after the first import too."""
    store.save_skill(tmp_path / "perfil", _skill())
    respuesta = cliente_web.get("/perfil")
    assert b'href="/perfil/importar"' in respuesta.data
    assert "¿Tienes otro CV que añadir?".encode("utf-8") in respuesta.data


def test_importar_cv_no_esta_en_la_cabecera(cliente_web):
    """Only lives on My profile, not in the global navigation — with
    Templates and Support already there, the header was getting too crowded."""
    respuesta = cliente_web.get("/ajustes")
    assert b'href="/perfil/importar"' not in respuesta.data


def test_subir_un_formato_no_soportado_muestra_el_error(cliente_web):
    datos = {"fichero": (io.BytesIO(b"contenido cualquiera"), "cv.txt")}
    respuesta = cliente_web.post(
        "/perfil/importar", data=datos, content_type="multipart/form-data"
    )
    assert respuesta.status_code == 200
    assert "no es un formato soportado".encode("utf-8") in respuesta.data


# --------------------------------------------------------------------------
# End to end: upload -> extract -> analyze -> review -> save.
#
# The tests above check each piece on its own (extraction, analysis with a
# fake ClienteIA, saving what was checked); this is the only one that shows
# the VIEW wires them together correctly. `create_client` is swapped for a
# double (touching neither the network nor a real key) so the whole chain
# can be tested without spending a real call.
# --------------------------------------------------------------------------


class _ClienteFalsoDisponible:
    """Always "available" and always returns the same proposal — there is
    no need to vary the response to show the view knows how to use it."""

    def __init__(self, respuesta: str):
        self.respuesta = respuesta

    def complete(self, sistema: str, usuario: str) -> str:
        return self.respuesta

    def available(self) -> bool:
        return True


def _respuesta_ia() -> str:
    import json

    return json.dumps(
        {
            "experiencias": [
                {
                    "titulo": {"es": "Ingeniera de Datos", "en": "Data Engineer"},
                    "periodo": {"es": "2025 - actualidad", "en": "2025 - present"},
                    "bullets": {
                        "es": ["Pipeline de ingesta con Airflow"],
                        "en": ["Ingestion pipeline with Airflow"],
                    },
                    "stack": {"es": "Python, Airflow", "en": "Python, Airflow"},
                    "keywords": ["airflow", "etl"],
                }
            ],
            "skills": [
                {"nombre": {"es": "SQL", "en": "SQL"}, "categoria": "dato", "keywords": ["sql"]}
            ],
            "skills_personales": [
                {
                    "nombre": {"es": "Trabajo en equipo", "en": "Teamwork"},
                    "keywords": ["team player"],
                }
            ],
            "idiomas": [
                {
                    "nombre": {"es": "Inglés", "en": "English"},
                    "nivel": {"es": "C1 Avanzado", "en": "C1 Advanced"},
                    "keywords": ["advanced english"],
                }
            ],
        },
        ensure_ascii=False,
    )


def _docx_de_prueba() -> bytes:
    import docx

    documento = docx.Document()
    documento.add_paragraph("Ana Ejemplo — Ingeniera de Datos")
    documento.add_paragraph("2025 - actualidad: pipeline de ingesta con Airflow y Python.")
    buffer = io.BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def test_de_punta_a_punta_subir_analizar_revisar_y_guardar(cliente_web, tmp_path, monkeypatch):
    import ancla.web.views.import_cv as vista_importar

    monkeypatch.setattr(
        vista_importar,
        "create_client",
        lambda proveedor, clave, url_base="", modelo="": _ClienteFalsoDisponible(_respuesta_ia()),
    )

    # 1) Upload a real .docx.
    subida = cliente_web.post(
        "/perfil/importar",
        data={"fichero": (io.BytesIO(_docx_de_prueba()), "mi-cv.docx")},
        content_type="multipart/form-data",
    )
    assert subida.status_code == 302
    assert subida.location.endswith("/perfil/importar/revisar")

    # 2) The review screen shows what the analysis returned, not a
    #    placeholder or another test's data.
    revision = cliente_web.get("/perfil/importar/revisar")
    html = revision.data.decode("utf-8")
    assert "Ingeniera de Datos" in html
    assert "Airflow" in html
    assert "SQL" in html
    assert "Trabajo en equipo" in html
    assert "Inglés" in html

    # 3) Save the experience and the personal skill; discard the technical
    #    skill and the language, to exercise both paths (save/discard)
    #    across all four categories at once.
    save = cliente_web.post(
        "/perfil/importar/guardar",
        data={
            "exp-0": "1",
            "exp-0-titulo_es": "Ingeniera de Datos",
            "exp-0-titulo_en": "Data Engineer",
            "exp-0-periodo_es": "2025 - actualidad",
            "exp-0-periodo_en": "2025 - present",
            "exp-0-bullets_es": "Pipeline de ingesta con Airflow",
            "exp-0-bullets_en": "Ingestion pipeline with Airflow",
            "exp-0-stack_es": "Python, Airflow",
            "exp-0-stack_en": "Python, Airflow",
            # "skill-0" is not sent: it stays discarded.
            "skillpersonal-0": "1",
            "skillpersonal-0-nombre_es": "Trabajo en equipo",
            "skillpersonal-0-nombre_en": "Teamwork",
            # "idioma-0" is not sent: it stays discarded.
        },
    )
    assert save.status_code == 302

    # 4) The real profile has the experience and the personal skill, and
    #    has NEITHER the discarded technical skill NOR the discarded language.
    perfil = store.load_profile(tmp_path / "perfil")
    experiencias = [e for e in perfil.experiences if e.title["es"] == "Ingeniera de Datos"]
    assert len(experiencias) == 1
    assert experiencias[0].bullets["es"] == ["Pipeline de ingesta con Airflow"]
    assert perfil.skill("sql") is None
    assert perfil.personal_skill("trabajo-en-equipo") is not None
    assert perfil.skill("trabajo-en-equipo") is None  # nunca en el catálogo técnico
    assert perfil.language("ingles") is None

    # 5) The pending import draft is cleared: the same batch cannot be
    #    "reviewed" again, and nothing is left dangling in the profile.
    assert modulo_importacion.load_import(tmp_path / "perfil") is None
