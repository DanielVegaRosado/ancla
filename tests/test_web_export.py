"""HTTP tests for exporting the proposal to a filled `.docx` (Camino A):
the download itself, the overflow question when the template has less room
than the proposal has experiences, and that a new template just needs its
two files dropped in `docx-templates/` to show up — no code change."""
from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document

from ancla.archive import repository as archivo
from ancla.profile import store
from ancla.profile.model import (
    AboutMe,
    Bilingual,
    CVStatus,
    Experience,
    SavedCV,
    Skill,
    SelectedAboutMe,
    SelectedExperience,
    Proposal,
    SpokenLanguage,
)
from ancla.web import create_app
from ancla.web import draft as modulo_borrador

PLANTILLAS_PRUEBA = Path(__file__).resolve().parent / "fixtures"

_EXPERIENCIAS = [
    ("exp-1", "Rol Uno · Empresa A", "Motivo A"),
    ("exp-2", "Rol Dos · Empresa B", "Motivo B"),
    ("exp-3", "Rol Tres · Empresa C", "Motivo C"),
]


def _perfil_en_disco(root: Path) -> None:
    for id_, titulo, _motivo in _EXPERIENCIAS:
        store.save_experience(
            root,
            Experience(
                id=id_,
                title=Bilingual(es=titulo, en=titulo),
                period=Bilingual(es="2023", en="2023"),
                bullets=Bilingual(es=["Bullet"], en=["Bullet"]),
                stack=Bilingual(es="Python", en="Python"),
            ),
        )
    store.save_skill(root, Skill(id="python", name=Bilingual(es="Python", en="Python")))
    store.save_about_me(
        root, AboutMe(template=Bilingual(es="Trabajo con {GROUP_A_1}.", en="I work with {GROUP_A_1}."))
    )
    store.save_personal_skill(root, Skill(id="equipo", name=Bilingual(es="Trabajo en equipo", en="Teamwork")))
    store.save_language(
        root,
        SpokenLanguage(
            id="en", name=Bilingual(es="Inglés", en="English"), level=Bilingual(es="C1", en="C1")
        ),
    )


def _propuesta(cuantas: int) -> Proposal:
    return Proposal(
        language="es",
        about_me=SelectedAboutMe(group_a=[], group_b=[], text="Sobre mí de la propuesta.", reason=""),
        skills=["python"],
        experiences=[
            SelectedExperience(id=id_, reason=motivo) for id_, _titulo, motivo in _EXPERIENCIAS[:cuantas]
        ],
    )


def _cliente(tmp_path: Path, n_experiencias: int, plantillas_root: Path = PLANTILLAS_PRUEBA):
    root = tmp_path / "perfil"
    _perfil_en_disco(root)
    app = create_app(
        raiz_perfil=root,
        settings_path=tmp_path / "ajustes.json",
        docx_templates_root=plantillas_root,
    )
    app.config["TESTING"] = True
    modulo_borrador.save_draft(
        root,
        modulo_borrador.Draft(
            vacante="Buscamos Python.", empresa="ACME", puesto="Backend", propuesta=_propuesta(n_experiencias)
        ),
    )
    return app.test_client()


def _cliente_con_cv(tmp_path: Path, n_experiencias: int, plantillas_root: Path = PLANTILLAS_PRUEBA) -> tuple:
    root = tmp_path / "perfil"
    _perfil_en_disco(root)
    app = create_app(
        raiz_perfil=root,
        settings_path=tmp_path / "ajustes.json",
        docx_templates_root=plantillas_root,
    )
    app.config["TESTING"] = True
    cv = SavedCV(
        id="acme-backend-2026-08-13",
        date=date(2026, 8, 13),
        company="ACME",
        position="Backend",
        posting="Buscamos Python.",
        proposal=_propuesta(n_experiencias),
        status=CVStatus.DRAFT,
    )
    archivo.save(root, cv)
    return app.test_client(), cv.id


def _filas_experiencia(contenido: bytes) -> list:
    documento = Document(BytesIO(contenido))
    tabla_principal = documento.tables[0].rows[0].cells[1]
    return tabla_principal.tables[0].rows


def test_descargar_sin_desbordamiento_devuelve_un_docx(tmp_path: Path):
    cliente = _cliente(tmp_path, n_experiencias=2)  # cabe justo en la plantilla de prueba (capacidad 2)

    respuesta = cliente.post("/propuesta/exportar", data={"plantilla_id": "prueba"})

    assert respuesta.status_code == 200
    assert respuesta.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "cv-prueba.docx" in respuesta.headers["Content-Disposition"]
    filas = _filas_experiencia(respuesta.data)
    assert len(filas) == 2


def test_menos_experiencias_seleccionadas_que_capacidad_se_completan_con_el_perfil(tmp_path: Path):
    """The test template has capacity 2; the proposal only chose 1. Instead
    of exporting with a gap, it's topped up with another profile experience
    that wasn't in the selection (no relevance reason attached)."""
    cliente = _cliente(tmp_path, n_experiencias=1)

    respuesta = cliente.post("/propuesta/exportar", data={"plantilla_id": "prueba"})

    assert respuesta.status_code == 200
    filas = _filas_experiencia(respuesta.data)
    assert len(filas) == 2
    assert "Rol Uno" in filas[0].cells[0].text
    assert "Rol Dos" in filas[1].cells[0].text


def test_mas_experiencias_que_capacidad_pregunta_por_el_desbordamiento(tmp_path: Path):
    cliente = _cliente(tmp_path, n_experiencias=3)  # la plantilla de prueba solo tiene sitio para 2

    respuesta = cliente.post("/propuesta/exportar", data={"plantilla_id": "prueba"})

    assert respuesta.status_code == 200
    assert respuesta.mimetype == "text/html"
    html = respuesta.data.decode("utf-8")
    assert "Rol Tres · Empresa C" in html  # la que se quedaría fuera
    assert "Motivo C" in html  # y su motivo


def test_desbordamiento_relevantes_recorta_a_la_capacidad(tmp_path: Path):
    cliente = _cliente(tmp_path, n_experiencias=3)

    respuesta = cliente.post(
        "/propuesta/exportar", data={"plantilla_id": "prueba", "desbordamiento": "relevantes"}
    )

    assert respuesta.status_code == 200
    filas = _filas_experiencia(respuesta.data)
    assert len(filas) == 2
    assert "Rol Uno" in filas[0].cells[0].text
    assert "Rol Dos" in filas[1].cells[0].text


def test_desbordamiento_todas_no_recorta_nada(tmp_path: Path):
    cliente = _cliente(tmp_path, n_experiencias=3)

    respuesta = cliente.post("/propuesta/exportar", data={"plantilla_id": "prueba", "desbordamiento": "todas"})

    assert respuesta.status_code == 200
    filas = _filas_experiencia(respuesta.data)
    assert len(filas) == 3
    assert "Rol Tres" in filas[2].cells[0].text


def test_plantilla_desconocida_redirige_con_aviso(tmp_path: Path):
    cliente = _cliente(tmp_path, n_experiencias=1)
    respuesta = cliente.post("/propuesta/exportar", data={"plantilla_id": "no-existe"}, follow_redirects=True)
    assert respuesta.status_code == 200
    assert "ya no está disponible" in respuesta.data.decode("utf-8")


def test_anadir_una_plantilla_nueva_la_hace_aparecer_en_el_desplegable(tmp_path: Path):
    """Demonstrates the acceptance criterion: registering a template
    doesn't require touching code, just dropping the `.docx` with its tags
    and its `.yaml`."""
    plantillas = tmp_path / "mis-plantillas"
    plantillas.mkdir()
    (plantillas / "corporativa.docx").write_bytes((PLANTILLAS_PRUEBA / "prueba.docx").read_bytes())
    (plantillas / "corporativa.yaml").write_text(
        "nombre: Corporativa Clásica\ncapacidad_experiencias: 4\n", encoding="utf-8"
    )

    cliente = _cliente(tmp_path, n_experiencias=1, plantillas_root=plantillas)
    respuesta = cliente.get("/propuesta")

    assert respuesta.status_code == 200
    html = respuesta.data.decode("utf-8")
    assert "Corporativa Clásica" in html
    assert '<option value="corporativa"' in html


def test_el_desplegable_muestra_el_nombre_en_el_idioma_de_la_propuesta(tmp_path: Path):
    """The dropdown text sits next to other content already shown in the
    proposal's own language (`propuesta.language`), not the interface
    language — the same field a bilingual template name resolves against."""
    plantillas = tmp_path / "mis-plantillas"
    plantillas.mkdir()
    (plantillas / "corporativa.docx").write_bytes((PLANTILLAS_PRUEBA / "prueba.docx").read_bytes())
    (plantillas / "corporativa.yaml").write_text(
        "nombre:\n  es: Corporativa Clásica\n  en: Classic Corporate\ncapacidad_experiencias: 4\n",
        encoding="utf-8",
    )

    root = tmp_path / "perfil"
    _perfil_en_disco(root)
    app = create_app(raiz_perfil=root, settings_path=tmp_path / "ajustes.json", docx_templates_root=plantillas)
    app.config["TESTING"] = True
    modulo_borrador.save_draft(
        root,
        modulo_borrador.Draft(
            vacante="We're hiring.",
            empresa="ACME",
            puesto="Backend",
            propuesta=Proposal(
                language="en",
                about_me=SelectedAboutMe(group_a=[], group_b=[], text="About me.", reason=""),
                skills=["python"],
                experiences=[SelectedExperience(id="exp-1", reason="Motivo A")],
            ),
        ),
    )

    html = app.test_client().get("/propuesta").data.decode("utf-8")
    assert "Classic Corporate" in html
    assert "Corporativa Clásica" not in html


def test_sin_plantillas_el_desplegable_no_aparece(tmp_path: Path):
    cliente = _cliente(tmp_path, n_experiencias=1, plantillas_root=tmp_path / "vacio")
    respuesta = cliente.get("/propuesta")
    assert respuesta.status_code == 200
    assert b"<select" not in respuesta.data or b'name="plantilla_id"' not in respuesta.data


# --------------------------------------------------------------------------
# Exporting a CV already saved to the archive (not just the current draft)
# --------------------------------------------------------------------------


def test_la_pantalla_de_un_cv_guardado_ofrece_el_desplegable_de_plantillas(tmp_path: Path):
    cliente, id_ = _cliente_con_cv(tmp_path, n_experiencias=1)
    respuesta = cliente.get(f"/cvs/{id_}")
    assert respuesta.status_code == 200
    html = respuesta.data.decode("utf-8")
    assert "Descargar CV maquetado" in html
    assert f'action="/cvs/{id_}/exportar"' in html


def test_descargar_un_cv_guardado_devuelve_un_docx(tmp_path: Path):
    cliente, id_ = _cliente_con_cv(tmp_path, n_experiencias=2)

    respuesta = cliente.post(f"/cvs/{id_}/exportar", data={"plantilla_id": "prueba"})

    assert respuesta.status_code == 200
    assert respuesta.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert f"{id_}-prueba.docx" in respuesta.headers["Content-Disposition"]
    assert len(_filas_experiencia(respuesta.data)) == 2


def test_descargar_un_cv_guardado_con_desbordamiento_pregunta_y_respeta_la_eleccion(tmp_path: Path):
    cliente, id_ = _cliente_con_cv(tmp_path, n_experiencias=3)

    pregunta = cliente.post(f"/cvs/{id_}/exportar", data={"plantilla_id": "prueba"})
    assert pregunta.mimetype == "text/html"
    html = pregunta.data.decode("utf-8")
    assert "Rol Tres · Empresa C" in html
    assert f'action="/cvs/{id_}/exportar"' in html  # el formulario reenvía a la misma ruta de CV

    respuesta = cliente.post(f"/cvs/{id_}/exportar", data={"plantilla_id": "prueba", "desbordamiento": "todas"})
    assert len(_filas_experiencia(respuesta.data)) == 3


def test_exportar_un_cv_que_no_existe_avisa_y_redirige(tmp_path: Path):
    cliente, _id = _cliente_con_cv(tmp_path, n_experiencias=1)
    respuesta = cliente.post("/cvs/no-existe/exportar", data={"plantilla_id": "prueba"}, follow_redirects=True)
    assert respuesta.status_code == 200
    assert "No se encuentra el CV" in respuesta.data.decode("utf-8")
