"""Tests for `ancla/export/`: discovering `.docx` templates from their YAML
sidecars, and building the fill context / rendering a real `.docx`."""
from __future__ import annotations

import dataclasses
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from docx.shared import Pt

from ancla.export import fill, templates
from ancla.profile.model import (
    Bilingual,
    Education,
    Experience,
    Profile,
    Proposal,
    SelectedAboutMe,
    SelectedExperience,
    Skill,
    SpokenLanguage,
)

PLANTILLA_PRUEBA = Path(__file__).resolve().parent / "fixtures" / "prueba.docx"


def _perfil() -> Profile:
    experiencias = [
        Experience(
            id="proyecto-a",
            title=Bilingual(es="Ingeniero de Datos · ACME", en="Data Engineer · ACME"),
            period=Bilingual(es="2023 — 2024", en="2023 — 2024"),
            bullets=Bilingual(es=["Bullet A1", "Bullet A2"], en=["Bullet A1 EN", "Bullet A2 EN"]),
            stack=Bilingual(es="Python", en="Python"),
        ),
        Experience(
            id="proyecto-b",
            title=Bilingual(es="Backend Developer · Nubelia", en="Backend Developer · Nubelia"),
            period=Bilingual(es="2021 — 2023", en="2021 — 2023"),
            bullets=Bilingual(es=["Bullet B1"], en=["Bullet B1 EN"]),
            stack=Bilingual(es="Go", en="Go"),
        ),
    ]
    skills = [Skill(id="python", name=Bilingual(es="Python", en="Python"))]
    personales = [Skill(id="equipo", name=Bilingual(es="Trabajo en equipo", en="Teamwork"))]
    idiomas = [SpokenLanguage(id="en", name=Bilingual(es="Inglés", en="English"), level=Bilingual(es="C1", en="C1"))]
    educacion = [
        Education(
            id="grado",
            title=Bilingual(es="Grado en Ingeniería Informática", en="BSc in Computer Engineering"),
            institution=Bilingual(es="UEMC", en="UEMC"),
            period=Bilingual(es="2023 — 2027", en="2023 — 2027"),
        )
    ]
    return Profile(
        experiences=experiencias,
        skills=skills,
        personal_skills=personales,
        languages=idiomas,
        education=educacion,
        contact=["+34 000 000 000", "tu-email@ejemplo.com"],
        headline=Bilingual(es="Ingeniero Informático", en="Computer Engineer"),
    )


def _propuesta() -> Proposal:
    return Proposal(
        language="es",
        about_me=SelectedAboutMe(group_a=[], group_b=[], text="Sobre mí de prueba.", reason=""),
        skills=["python"],
        experiences=[
            SelectedExperience(id="proyecto-a", reason="Encaja con la vacante."),
            SelectedExperience(id="proyecto-b", reason="Completa la sección."),
        ],
    )


# --------------------------------------------------------------------------
# templates.py
# --------------------------------------------------------------------------


def test_lista_plantillas_lee_la_plantilla_de_prueba_de_fixtures():
    plantillas = templates.list_templates(PLANTILLA_PRUEBA.parent)
    ids = [p.id for p in plantillas]
    assert "prueba" in ids
    prueba = next(p for p in plantillas if p.id == "prueba")
    assert prueba.name.es == "Plantilla de prueba"
    assert prueba.name.en == "Plantilla de prueba"
    assert prueba.capacity_experiences == 2
    # No declara `geometria`: no participa del ajuste de densidad, y eso
    # no debe romper nada al listarla.
    assert prueba.geometry is None


def test_lee_la_geometria_declarada_en_el_yaml(tmp_path: Path):
    (tmp_path / "densa.docx").write_bytes(b"contenido falso")
    (tmp_path / "densa.yaml").write_text(
        "nombre: Densa\n"
        "capacidad_experiencias: 3\n"
        "geometria:\n"
        "  ancho_columna: 300.0\n"
        "  ancho_columna_titulo: 280.0\n"
        "  alto_pagina: 841.92\n"
        "  margen_inferior_seguridad: 35.0\n"
        "  altura_hasta_experiencia: 251.0\n"
        "  altura_tras_experiencia: 0.0\n"
        "  tamano_nombre_primero: 22.5\n"
        "  tamano_nombre_resto: 0.0\n"
        "  tamano_titular: 10.5\n"
        "  tamano_texto_antes_experiencia: 10.5\n"
        "  linea_texto_antes_experiencia: 1.35\n"
        "  tamano_texto_tras_experiencia: 0.0\n"
        "  linea_texto_tras_experiencia: 0.0\n"
        "  tamano_titulo: 10.5\n"
        "  tamano_cuerpo: 9.5\n"
        "  tamano_stack: 9.0\n"
        "  espacio_titulo: 8.0\n"
        "  espacio_bullet: 1.5\n"
        "  linea_bullet: 1.1\n"
        "  calibracion: 1.12\n"
        "  fuente: Aileron\n",
        encoding="utf-8",
    )
    plantilla = templates.find_template(tmp_path, "densa")
    assert plantilla.geometry == templates.TemplateGeometry(
        ancho_columna=300.0,
        ancho_columna_titulo=280.0,
        alto_pagina=841.92,
        margen_inferior_seguridad=35.0,
        altura_hasta_experiencia=251.0,
        altura_tras_experiencia=0.0,
        tamano_nombre_primero=22.5,
        tamano_nombre_resto=0.0,
        tamano_titular=10.5,
        tamano_texto_antes_experiencia=10.5,
        linea_texto_antes_experiencia=1.35,
        tamano_texto_tras_experiencia=0.0,
        linea_texto_tras_experiencia=0.0,
        tamano_titulo=10.5,
        tamano_cuerpo=9.5,
        tamano_stack=9.0,
        espacio_titulo=8.0,
        espacio_bullet=1.5,
        linea_bullet=1.1,
        calibracion=1.12,
        fuente="Aileron",
    )


def test_una_geometria_incompleta_se_ignora_en_vez_de_romper(tmp_path: Path):
    """Same rule as a broken `.yaml` overall: a template someone is still
    preparing shouldn't crash the Proposal screen for everyone else's."""
    (tmp_path / "a-medias.docx").write_bytes(b"contenido falso")
    (tmp_path / "a-medias.yaml").write_text(
        "nombre: A medias\ncapacidad_experiencias: 3\ngeometria:\n  ancho_columna: 300.0\n",
        encoding="utf-8",
    )
    plantilla = templates.find_template(tmp_path, "a-medias")
    assert plantilla is not None
    assert plantilla.geometry is None


def test_carpeta_inexistente_no_da_error(tmp_path: Path):
    assert templates.list_templates(tmp_path / "no-existe") == []


def test_un_docx_sin_yaml_hermano_se_ignora(tmp_path: Path):
    (tmp_path / "huerfano.docx").write_bytes(b"contenido falso")
    assert templates.list_templates(tmp_path) == []


def test_un_yaml_con_formato_invalido_se_ignora(tmp_path: Path):
    (tmp_path / "rota.docx").write_bytes(b"contenido falso")
    (tmp_path / "rota.yaml").write_text("esto: [no cierra", encoding="utf-8")
    assert templates.list_templates(tmp_path) == []


def test_find_template_devuelve_none_si_no_existe(tmp_path: Path):
    assert templates.find_template(tmp_path, "no-existe") is None


def test_anadir_una_plantilla_nueva_no_exige_tocar_codigo(tmp_path: Path):
    """The entire circuit for registering a template is dropping the two
    files in the folder: no lists or name-based `if`s in the code."""
    (tmp_path / "corporativa.docx").write_bytes(b"contenido falso")
    (tmp_path / "corporativa.yaml").write_text(
        "nombre:\n  es: Corporativa\n  en: Corporate\ncapacidad_experiencias: 5\n", encoding="utf-8"
    )
    plantillas = templates.list_templates(tmp_path)
    assert len(plantillas) == 1
    assert plantillas[0].name.es == "Corporativa"
    assert plantillas[0].name.en == "Corporate"
    assert plantillas[0].capacity_experiences == 5


# --------------------------------------------------------------------------
# fill.py
# --------------------------------------------------------------------------


def test_resolved_selection_descarta_experiencias_que_ya_no_existen():
    perfil = _perfil()
    propuesta = _propuesta()
    propuesta.experiences.append(SelectedExperience(id="no-existe", reason=""))
    seleccion = fill.resolved_selection(propuesta, perfil)
    assert [s.id for s, _ in seleccion] == ["proyecto-a", "proyecto-b"]


def test_build_context_expone_el_catalogo_de_campos():
    perfil = _perfil()
    propuesta = _propuesta()
    experiencias = fill.resolved_experiences(propuesta, perfil)
    contexto = fill.build_context(propuesta, perfil, experiencias, "Daniel Vega")

    assert contexto["nombre"] == "Daniel Vega"
    assert contexto["sobre_mi"] == "Sobre mí de prueba."
    assert contexto["skills"] == ["Python"]
    assert contexto["idiomas"] == ["Inglés — C1"]
    assert contexto["skills_personales"] == ["Trabajo en equipo"]
    assert len(contexto["experiencias"]) == 2
    primera = contexto["experiencias"][0]
    assert primera["puesto"] == "Ingeniero de Datos · ACME"
    assert primera["empresa"] == ""
    assert primera["fechas"] == "2023 — 2024"
    assert primera["bullets"] == ["Bullet A1", "Bullet A2"]
    assert primera["stack"] == "Python"
    assert contexto["contacto"] == ["+34 000 000 000", "tu-email@ejemplo.com"]
    assert contexto["titular"] == "Ingeniero Informático"
    assert contexto["educacion"] == [
        {"titulo": "Grado en Ingeniería Informática", "centro": "UEMC", "fechas": "2023 — 2027"}
    ]
    assert contexto["foto"] == ""


def test_nunca_reescribe_los_bullets_del_usuario():
    """Product rule 2: bullets are shown verbatim, letter for letter — the
    context doesn't touch them, only copies them from the experience."""
    perfil = _perfil()
    propuesta = _propuesta()
    experiencias = fill.resolved_experiences(propuesta, perfil)
    contexto = fill.build_context(propuesta, perfil, experiencias, "")
    assert contexto["experiencias"][0]["bullets"] == perfil.experience("proyecto-a").bullets["es"]


def test_sustituye_el_guion_no_separable_por_uno_normal():
    """Narrow, documented exception to rule 2: several export fonts don't
    ship the U+2011 glyph and Word falls back to another font just for that
    character. This is a typographic-compatibility fix, not a content
    rewrite — visually identical, and the saved profile is never touched
    (it only happens in `build_context`, on the way to the `.docx`)."""
    guion_no_separable = "H‑Z‑H"
    perfil = dataclasses.replace(
        _perfil(), headline=Bilingual(es=f"Rol {guion_no_separable}", en=f"Role {guion_no_separable}")
    )
    perfil.experience("proyecto-a").bullets["es"][0] = f"Secuencia {guion_no_separable}"
    propuesta = _propuesta()
    experiencias = fill.resolved_experiences(propuesta, perfil)
    contexto = fill.build_context(propuesta, perfil, experiencias, "")

    assert contexto["experiencias"][0]["bullets"][0] == "Secuencia H-Z-H"
    assert contexto["titular"] == "Rol H-Z-H"
    assert "‑" not in contexto["experiencias"][0]["bullets"][0]


def _experiencia_de_prueba(n_bullets: int = 2, stack: str = "Python") -> dict:
    return {
        "puesto": "Data Engineer",
        "bullets": ["Bullet corto." for _ in range(n_bullets)],
        "stack": stack,
    }


def _geometria_de_prueba() -> templates.TemplateGeometry:
    """Mirrors `corporativa-clasica.yaml`'s `geometria` block — kept as a
    literal copy on purpose, so a regression in what the build script
    writes would show up here as a mismatch, not get silently inherited."""
    return templates.TemplateGeometry(
        ancho_columna=306.8352,
        ancho_columna_titulo=230.0352,
        alto_pagina=841.92,
        margen_inferior_seguridad=35.0,
        altura_hasta_experiencia=115.5,
        altura_tras_experiencia=0.0,
        tamano_nombre_primero=22.5,
        tamano_nombre_resto=0.0,
        tamano_titular=10.5,
        tamano_texto_antes_experiencia=10.5,
        linea_texto_antes_experiencia=1.35,
        tamano_texto_tras_experiencia=0.0,
        linea_texto_tras_experiencia=0.0,
        tamano_titulo=10.5,
        tamano_cuerpo=9.5,
        tamano_stack=9.0,
        espacio_titulo=8.0,
        espacio_bullet=1.5,
        linea_bullet=1.1,
        calibracion=1.12,
        fuente="Aileron",
    )


def _contexto_de_prueba(
    *, experiencias=None, sobre_mi="", nombre_primero="", nombre_resto="", titular=""
) -> dict:
    return {
        "experiencias": experiencias or [],
        "sobre_mi": sobre_mi,
        "nombre_primero": nombre_primero,
        "nombre_resto": nombre_resto,
        "titular": titular,
    }


def test_lineas_estimadas_cuenta_el_envoltorio_por_ancho_real():
    """Validated against 10 real bullets rendered through Word (10/10
    matched)."""
    geometria = _geometria_de_prueba()
    texto_una_linea = "Record multi-step tool calls, observations, and outputs."
    texto_dos_lineas = (
        "Open-source Flask app that adapts a CV to each job posting by "
        "selecting from verified facts, never inventing content."
    )
    anchos = fill.metricas_aileron.ANCHOS_REGULAR
    assert fill._lineas_estimadas(texto_una_linea, anchos, geometria.tamano_cuerpo, geometria.ancho_columna) == 1
    assert fill._lineas_estimadas(texto_dos_lineas, anchos, geometria.tamano_cuerpo, geometria.ancho_columna) == 2
    assert fill._lineas_estimadas("", anchos, geometria.tamano_cuerpo, geometria.ancho_columna) == 0


def test_altura_bloque_usa_la_tabla_de_anchos_de_la_fuente_declarada():
    """Antes de esto, `_altura_bloque` siempre medía con las tablas de
    Aileron, sin importar qué fuente declarase la plantilla — Montserrat es
    sensiblemente más ancha, así que un mismo bullet puede envolver a una
    línea de más con Montserrat sin que Aileron lo prediga, si la tabla no
    cambia según la plantilla."""
    # A esta anchura de columna, este texto envuelve a 1 línea medido con
    # Aileron y a 2 con Montserrat — comprobado aparte contra las dos tablas.
    experiencia = _experiencia_de_prueba(n_bullets=0)
    experiencia["bullets"] = ["Optimized hyperparameters iteratively using cross validation folds today"]
    geometria_aileron = _geometria_de_prueba()
    geometria_montserrat = dataclasses.replace(geometria_aileron, fuente="Montserrat")
    assert fill._altura_bloque(experiencia, 1.0, geometria_aileron) < fill._altura_bloque(
        experiencia, 1.0, geometria_montserrat
    )


def test_altura_bloque_crece_con_la_escala():
    experiencia = _experiencia_de_prueba()
    geometria = _geometria_de_prueba()
    assert fill._altura_bloque(experiencia, 1.5, geometria) > fill._altura_bloque(experiencia, 1.0, geometria)


def test_escala_necesaria_es_1_si_el_contenido_ya_llena_la_pagina():
    """With many long experiences, the block already reaches the margin on
    its own — nothing meaningful to scale in either direction."""
    muchas = [_experiencia_de_prueba(n_bullets=4) for _ in range(8)]
    assert fill._escala_necesaria(muchas, _geometria_de_prueba(), _contexto_de_prueba()) == pytest.approx(1.0, abs=0.05)


def test_escala_necesaria_crece_cuando_hay_poco_contenido():
    geometria = _geometria_de_prueba()
    contexto = _contexto_de_prueba()
    poca = [_experiencia_de_prueba(n_bullets=1)]
    mucha = [_experiencia_de_prueba(n_bullets=4) for _ in range(8)]
    assert fill._escala_necesaria(poca, geometria, contexto) > fill._escala_necesaria(mucha, geometria, contexto) >= 0.9


def test_escala_necesaria_respeta_el_tope_maximo():
    """A single experience with one short bullet has nowhere near enough
    text to fill the page through scaling alone — it stops at the cap
    instead of growing without limit to force a fit."""
    casi_vacia = [_experiencia_de_prueba(n_bullets=1)]
    assert fill._escala_necesaria(casi_vacia, _geometria_de_prueba(), _contexto_de_prueba()) == fill._ESCALA_MAXIMA


def test_escala_necesaria_encoge_por_debajo_de_1_cuando_sobra_contenido():
    """Real bug this fixes: real bullets and a real "About me" can add up
    to more than the page has room for even at the template's own declared
    capacity — before this, `escala` could never go below 1.0, so instead
    of tightening a little to fit, the CV just spilled onto a second page.
    More overflowing content has to shrink further, not less."""
    geometria = _geometria_de_prueba()
    contexto = _contexto_de_prueba()
    bastante = [_experiencia_de_prueba(n_bullets=4) for _ in range(8)]
    demasiado = [_experiencia_de_prueba(n_bullets=4) for _ in range(9)]
    escala_bastante = fill._escala_necesaria(bastante, geometria, contexto)
    escala_demasiado = fill._escala_necesaria(demasiado, geometria, contexto)
    assert fill._ESCALA_MINIMA <= escala_demasiado < escala_bastante < 1.0


def test_escala_necesaria_reserva_mas_hueco_cuando_el_texto_final_es_mas_largo():
    """Real bug this fixes: a fixed guess for how much "About me" (or any
    text after the experience block) takes up either leaves a blank gap
    when the real text is short, or spills the CV to a second page when
    it's longer than guessed — measuring the actual text avoids both."""
    geometria = dataclasses.replace(
        _geometria_de_prueba(), tamano_texto_tras_experiencia=9.5, linea_texto_tras_experiencia=1.35
    )
    experiencias = [_experiencia_de_prueba(n_bullets=4) for _ in range(5)]
    texto_corto = "Breve."
    texto_largo = "Una frase mucho más larga que ocupa bastantes más líneas de la columna. " * 10
    assert fill._escala_necesaria(
        experiencias, geometria, _contexto_de_prueba(sobre_mi=texto_largo)
    ) < fill._escala_necesaria(experiencias, geometria, _contexto_de_prueba(sobre_mi=texto_corto))


def test_escala_necesaria_reserva_mas_hueco_cuando_el_rol_es_mas_largo():
    """Same bug, other end of the page: the name/role block above the
    experience section varies per profile too (a long job title can wrap)
    — a fixed guess for it would be just as wrong as a fixed guess for
    "About me" was."""
    geometria = _geometria_de_prueba()
    experiencias = [_experiencia_de_prueba(n_bullets=4) for _ in range(5)]
    rol_corto = "QA"
    rol_largo = "Senior Machine Learning Engineer and Data Platform Architect for Enterprise Systems"
    assert fill._escala_necesaria(
        experiencias, geometria, _contexto_de_prueba(titular=rol_largo)
    ) < fill._escala_necesaria(experiencias, geometria, _contexto_de_prueba(titular=rol_corto))


def test_altura_cabecera_combina_nombre_y_apellido_si_comparten_linea():
    """Corporativa Clásica pone nombre y apellido en la misma línea/tamaño
    (`tamano_nombre_resto=0`, la convención que usa `TemplateGeometry` para
    decir "no hay segunda línea"); Minimalista Cálida los apila en dos
    líneas de tamaño distinto. Medir cada uno como línea propia cuando en
    realidad comparten una sola doblaría la altura estimada."""
    geometria_combinada = _geometria_de_prueba()  # tamano_nombre_resto=0.0
    geometria_apilada = dataclasses.replace(geometria_combinada, tamano_nombre_resto=geometria_combinada.tamano_nombre_primero)
    contexto = _contexto_de_prueba(nombre_primero="Daniel", nombre_resto="Vega")
    assert fill._altura_cabecera(contexto, geometria_combinada) < fill._altura_cabecera(contexto, geometria_apilada)


def test_aflojar_espaciado_escala_un_run_sin_tamano_explicito_usando_tamano_base():
    """Real bug this fixes: `run.font.size` reads back as `None` — not the
    run's real size — whenever that size happens to match what the run
    already inherits from its paragraph/style, because Word treats the
    explicit value as redundant and strips it the next time the file is
    opened and saved (confirmed happening during font embedding in
    `herramientas/incrustar_fuentes.py`, for exactly the runs whose size
    equals the template's `Normal` style default). Before `tamano_base`,
    such a run was silently skipped instead of scaled — visually correct
    at `escala=1.0` by coincidence, silently wrong at any other escala."""
    documento = Document()
    p = documento.add_paragraph("bullet sin tamano explicito")
    assert p.runs[0].font.size is None  # a fresh run never has one unless set

    fill._aflojar_espaciado(p, escala=1.5, tamano_base=9.5)

    # OOXML's `w:sz` only has half-point granularity, so the exact value
    # rounds a little on the way in — a real constraint of the file format,
    # not a precision bug — hence the tolerance instead of exact equality.
    assert p.runs[0].font.size.pt == pytest.approx(9.5 * fill._factor_tamano(1.5), abs=0.25)


def _plantilla_falsa_con_bloque_de_experiencia(n_filas: int) -> SimpleNamespace:
    """A minimal `.docx` shaped the way `_ajustar_densidad_experiencias`
    expects: a layout table (row 0 = sidebar, 1 = main column) → a nested
    experience table, one row per experience, each with a bullet paragraph
    carrying explicit spacing and font size.

    Wrapped in a `SimpleNamespace` with a `docx` attribute because
    `_ajustar_densidad_experiencias` receives a real `DocxTemplate` in
    production, which exposes the document that way (`documento.docx`) —
    see `fill.py::_recortar_a_circulo` for why `get_docx()` cannot be used
    instead."""
    documento = Document()
    layout = documento.add_table(rows=1, cols=2)
    columna_principal = layout.rows[0].cells[1]
    tabla_experiencias = columna_principal.add_table(rows=n_filas, cols=1)
    for fila in tabla_experiencias.rows:
        parrafo = fila.cells[0].paragraphs[0]
        run = parrafo.add_run("• bullet")
        run.font.size = Pt(9.5)
        parrafo.paragraph_format.space_after = Pt(3)
        parrafo.paragraph_format.line_spacing = 1.1
    return SimpleNamespace(docx=documento)


def test_ajustar_densidad_afloja_espaciado_y_letra_cuando_falta_contenido():
    plantilla = _plantilla_falsa_con_bloque_de_experiencia(n_filas=1)
    parrafo_antes = plantilla.docx.tables[0].rows[0].cells[1].tables[0].rows[0].cells[0].paragraphs[0]
    espaciado_antes = parrafo_antes.paragraph_format.space_after
    tamano_antes = parrafo_antes.runs[0].font.size

    experiencias = [_experiencia_de_prueba(n_bullets=1)]  # very little content, forces loosening
    fill._ajustar_densidad_experiencias(plantilla, _contexto_de_prueba(experiencias=experiencias), _geometria_de_prueba())

    parrafo_despues = plantilla.docx.tables[0].rows[0].cells[1].tables[0].rows[0].cells[0].paragraphs[0]
    assert parrafo_despues.paragraph_format.space_after > espaciado_antes
    assert parrafo_despues.runs[0].font.size > tamano_antes


def test_ajustar_densidad_no_revienta_si_la_plantilla_no_tiene_esa_estructura():
    """`prueba.docx`, or any template without a nested experience table,
    must not break the render — there's just nothing to loosen."""
    plantilla = SimpleNamespace(docx=Document())
    plantilla.docx.add_paragraph("CV without a two-column layout table")
    experiencias = [_experiencia_de_prueba(n_bullets=1)]
    contexto = _contexto_de_prueba(experiencias=experiencias)
    fill._ajustar_densidad_experiencias(plantilla, contexto, _geometria_de_prueba())  # must not raise


def test_ajustar_densidad_no_hace_nada_sin_experiencias():
    plantilla = _plantilla_falsa_con_bloque_de_experiencia(n_filas=1)
    contexto = _contexto_de_prueba(experiencias=[])
    fill._ajustar_densidad_experiencias(plantilla, contexto, _geometria_de_prueba())  # must not raise or scale anything


def test_render_produce_un_docx_con_el_contenido_real(tmp_path: Path):
    perfil = _perfil()
    propuesta = _propuesta()
    plantilla = templates.find_template(PLANTILLA_PRUEBA.parent, "prueba")
    experiencias = fill.resolved_experiences(propuesta, perfil)

    contenido = fill.render(plantilla, propuesta, perfil, experiencias, "Daniel Vega", tmp_path)

    salida = tmp_path / "salida.docx"
    salida.write_bytes(contenido)
    documento = Document(salida)

    texto_completo = "\n".join(p.text for p in documento.paragraphs)
    assert "Daniel Vega" in texto_completo

    # The template keeps its layout table and its nested experience table
    # (docxtpl fills existing runs, it doesn't rebuild the document) —
    # this isn't a blank Word doc with pasted text.
    assert len(documento.tables) == 1
    tabla_principal = documento.tables[0].rows[0].cells[1]
    assert "Sobre mí de prueba." in tabla_principal.text
    assert len(tabla_principal.tables) == 1
    filas_experiencias = tabla_principal.tables[0].rows
    assert len(filas_experiencias) == 2  # one row per experience, no leftover for/endfor tags
    assert "Ingeniero de Datos · ACME" in filas_experiencias[0].cells[0].text
    assert "Backend Developer · Nubelia" in filas_experiencias[1].cells[0].text


def test_render_con_foto_embebe_la_imagen(tmp_path: Path):
    """`store.photo_path` is the only thing that decides whether there's a
    photo — `render` doesn't need to know anything about profiles with or
    without one on its own."""
    import base64

    from ancla.profile import store as profile_store

    perfil = _perfil()
    propuesta = _propuesta()
    plantilla = templates.find_template(PLANTILLA_PRUEBA.parent, "prueba")
    experiencias = fill.resolved_experiences(propuesta, perfil)
    # A real, minimal 1x1 PNG: InlineImage needs to read its dimensions
    # from the file itself, an arbitrary byte string won't do.
    png_1x1 = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
        "+A8AAQUBAScY42YAAAAASUVORK5CYII="
    )
    profile_store.save_photo(tmp_path, "foto.png", png_1x1)

    # The test template has no {{ foto }} tag, so this only checks that
    # render() doesn't crash while building a real InlineImage that the
    # template simply ignores.
    contenido = fill.render(plantilla, propuesta, perfil, experiencias, "Daniel Vega", tmp_path)
    assert contenido[:2] == b"PK"  # still a valid .docx (zip)
