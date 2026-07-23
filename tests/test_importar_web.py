"""Tests de la pantalla Importar: estado efímero de la importación y las
rutas HTTP de punta a punta (subir, revisar, guardar seleccionadas, descartar).
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest

from cv_adaptativo.perfil import almacen
from cv_adaptativo.perfil.modelo import Bilingue, Experiencia, IdiomaHablado, Skill
from cv_adaptativo.web import crear_app
from cv_adaptativo.web import importacion as modulo_importacion


@pytest.fixture
def cliente_web(tmp_path: Path):
    app = crear_app(raiz_perfil=tmp_path / "perfil", ruta_ajustes=tmp_path / "ajustes.json")
    app.config["TESTING"] = True
    return app.test_client()


def _experiencia(id: str = "ml-dev") -> Experiencia:
    return Experiencia(
        id=id,
        titulo=Bilingue(es="ML Developer", en="ML Developer"),
        periodo=Bilingue(es="2026", en="2026"),
        bullets=Bilingue(es=["Pipeline completo"], en=["Full pipeline"]),
        stack=Bilingue(es="Python", en="Python"),
        keywords=["ml"],
    )


def _skill(id: str = "python") -> Skill:
    return Skill(id=id, nombre=Bilingue(es="Python", en="Python"), categoria="lenguaje", keywords=["py"])


def _skill_personal(id: str = "equipo") -> Skill:
    return Skill(id=id, nombre=Bilingue(es="Trabajo en equipo", en="Teamwork"), keywords=["team player"])


def _idioma(id: str = "ingles") -> IdiomaHablado:
    return IdiomaHablado(
        id=id,
        nombre=Bilingue(es="Inglés", en="English"),
        nivel=Bilingue(es="C1", en="C1"),
        keywords=["advanced english"],
    )


# --------------------------------------------------------------------------
# web/importacion.py: estado efímero, ida y vuelta
# --------------------------------------------------------------------------


def test_sin_importacion_guardada_no_hay_nada_que_cargar(tmp_path: Path):
    assert modulo_importacion.cargar_importacion(tmp_path) is None


def test_guardar_y_cargar_importacion_hace_ida_y_vuelta(tmp_path: Path):
    original = modulo_importacion.Importacion(
        experiencias=[_experiencia()], skills=[_skill()], avisos=["ojo con esto"]
    )
    modulo_importacion.guardar_importacion(tmp_path, original)
    recargada = modulo_importacion.cargar_importacion(tmp_path)
    assert recargada == original


def test_borrar_importacion_la_deja_indisponible(tmp_path: Path):
    modulo_importacion.guardar_importacion(
        tmp_path, modulo_importacion.Importacion(experiencias=[_experiencia()])
    )
    modulo_importacion.borrar_importacion(tmp_path)
    assert modulo_importacion.cargar_importacion(tmp_path) is None


def test_borrar_importacion_sin_fichero_no_falla(tmp_path: Path):
    modulo_importacion.borrar_importacion(tmp_path)  # no debe lanzar


def test_un_fichero_de_importacion_corrupto_no_revienta(tmp_path: Path):
    (tmp_path / modulo_importacion.NOMBRE_FICHERO).write_text("esto no es JSON", encoding="utf-8")
    assert modulo_importacion.cargar_importacion(tmp_path) is None


# --------------------------------------------------------------------------
# Rutas HTTP
# --------------------------------------------------------------------------


def test_ver_importar(cliente_web):
    respuesta = cliente_web.get("/perfil/importar")
    assert respuesta.status_code == 200
    assert "Importar".encode("utf-8") in respuesta.data


def test_importar_sin_fichero_ni_texto_pide_uno_de_los_dos(cliente_web):
    respuesta = cliente_web.post("/perfil/importar", data={"texto": ""})
    assert respuesta.status_code == 200
    assert "Sube un fichero o pega el texto".encode("utf-8") in respuesta.data


def test_importar_sin_clave_configurada_redirige_a_ajustes(cliente_web):
    respuesta = cliente_web.post(
        "/perfil/importar", data={"texto": "Un CV cualquiera con suficiente texto."}
    )
    assert respuesta.status_code == 302
    assert respuesta.location.endswith("/ajustes")


def test_revisar_sin_importacion_pendiente_redirige_a_importar(cliente_web):
    respuesta = cliente_web.get("/perfil/importar/revisar", follow_redirects=True)
    assert "No hay ninguna importación pendiente".encode("utf-8") in respuesta.data


def test_revisar_muestra_las_candidatas_guardadas(cliente_web, tmp_path: Path):
    modulo_importacion.guardar_importacion(
        tmp_path / "perfil",
        modulo_importacion.Importacion(experiencias=[_experiencia()], skills=[_skill()]),
    )
    respuesta = cliente_web.get("/perfil/importar/revisar")
    assert "ML Developer".encode("utf-8") in respuesta.data
    assert "Python".encode("utf-8") in respuesta.data


def test_guardar_solo_lo_marcado(cliente_web, tmp_path: Path):
    raiz = tmp_path / "perfil"
    modulo_importacion.guardar_importacion(
        raiz,
        modulo_importacion.Importacion(
            experiencias=[_experiencia("uno"), _experiencia("dos")], skills=[_skill()]
        ),
    )
    cliente_web.post(
        "/perfil/importar/guardar",
        data={
            "exp-0": "1",
            "exp-0-titulo_es": "ML Developer",
            "exp-0-titulo_en": "ML Developer",
            "exp-0-periodo_es": "2026",
            "exp-0-periodo_en": "2026",
            "exp-0-bullets_es": "Pipeline completo",
            "exp-0-bullets_en": "Full pipeline",
            "exp-0-stack_es": "Python",
            "exp-0-stack_en": "Python",
            # "exp-1" no viene en el formulario: no estaba marcado
        },
    )
    perfil = almacen.cargar_perfil(raiz)
    assert perfil.experiencia("uno") is not None
    assert perfil.experiencia("dos") is None
    assert perfil.skill("python") is None  # tampoco estaba marcado


def test_guardar_permite_editar_antes_de_confirmar(cliente_web, tmp_path: Path):
    raiz = tmp_path / "perfil"
    modulo_importacion.guardar_importacion(
        raiz, modulo_importacion.Importacion(skills=[_skill()])
    )
    cliente_web.post(
        "/perfil/importar/guardar",
        data={
            "skill-0": "1",
            "skill-0-nombre_es": "Python avanzado",
            "skill-0-nombre_en": "Advanced Python",
            "skill-0-categoria": "lenguaje",
        },
    )
    perfil = almacen.cargar_perfil(raiz)
    assert perfil.skill("python").nombre["es"] == "Python avanzado"


def test_una_candidata_incompleta_no_se_guarda_pero_avisa(cliente_web, tmp_path: Path):
    """Si al editar se borra el nombre en inglés, la validación normal la
    rechaza — igual que en el formulario manual — en vez de guardarla rota."""
    raiz = tmp_path / "perfil"
    modulo_importacion.guardar_importacion(
        raiz, modulo_importacion.Importacion(skills=[_skill()])
    )
    respuesta = cliente_web.post(
        "/perfil/importar/guardar",
        data={"skill-0": "1", "skill-0-nombre_es": "Python", "skill-0-nombre_en": ""},
        follow_redirects=True,
    )
    perfil = almacen.cargar_perfil(raiz)
    assert perfil.skill("python") is None
    assert "no se pudieron guardar".encode("utf-8") in respuesta.data


def test_guardar_una_skill_personal_importada_va_al_catalogo_correcto(cliente_web, tmp_path: Path):
    """Caso real que motivó esto: una sección "PERSONAL" del CV (Trabajo en
    equipo, Team player...) tiene que acabar en skills_personales, nunca
    mezclada con las skills técnicas."""
    raiz = tmp_path / "perfil"
    modulo_importacion.guardar_importacion(
        raiz, modulo_importacion.Importacion(skills_personales=[_skill_personal()])
    )
    cliente_web.post(
        "/perfil/importar/guardar",
        data={
            "skillpersonal-0": "1",
            "skillpersonal-0-nombre_es": "Trabajo en equipo",
            "skillpersonal-0-nombre_en": "Teamwork",
        },
    )
    perfil = almacen.cargar_perfil(raiz)
    assert perfil.skill_personal("equipo") is not None
    assert perfil.skill("equipo") is None  # nunca en el catálogo técnico


def test_guardar_un_idioma_importado(cliente_web, tmp_path: Path):
    raiz = tmp_path / "perfil"
    modulo_importacion.guardar_importacion(
        raiz, modulo_importacion.Importacion(idiomas=[_idioma()])
    )
    cliente_web.post(
        "/perfil/importar/guardar",
        data={
            "idioma-0": "1",
            "idioma-0-nombre_es": "Inglés",
            "idioma-0-nombre_en": "English",
            "idioma-0-nivel_es": "C1 Avanzado",
            "idioma-0-nivel_en": "C1 Advanced",
        },
    )
    perfil = almacen.cargar_perfil(raiz)
    assert perfil.idioma("ingles") is not None
    assert perfil.idioma("ingles").nivel["es"] == "C1 Avanzado"


def test_una_skill_personal_o_idioma_no_marcados_no_se_guardan(cliente_web, tmp_path: Path):
    raiz = tmp_path / "perfil"
    modulo_importacion.guardar_importacion(
        raiz,
        modulo_importacion.Importacion(
            skills_personales=[_skill_personal()], idiomas=[_idioma()]
        ),
    )
    cliente_web.post("/perfil/importar/guardar", data={})
    perfil = almacen.cargar_perfil(raiz)
    assert perfil.skill_personal("equipo") is None
    assert perfil.idioma("ingles") is None


def test_revisar_muestra_skills_personales_e_idiomas(cliente_web, tmp_path: Path):
    modulo_importacion.guardar_importacion(
        tmp_path / "perfil",
        modulo_importacion.Importacion(
            skills_personales=[_skill_personal()], idiomas=[_idioma()]
        ),
    )
    respuesta = cliente_web.get("/perfil/importar/revisar")
    html = respuesta.data.decode("utf-8")
    assert "Trabajo en equipo" in html
    assert "Inglés" in html and "C1" in html


def test_guardar_limpia_la_importacion_pendiente(cliente_web, tmp_path: Path):
    raiz = tmp_path / "perfil"
    modulo_importacion.guardar_importacion(raiz, modulo_importacion.Importacion(skills=[_skill()]))
    cliente_web.post("/perfil/importar/guardar", data={"skill-0": "1", "skill-0-nombre_es": "Python", "skill-0-nombre_en": "Python"})
    assert modulo_importacion.cargar_importacion(raiz) is None


def test_guardar_sin_importacion_pendiente_redirige(cliente_web):
    respuesta = cliente_web.post("/perfil/importar/guardar", data={})
    assert respuesta.status_code == 302


def test_descartar_borra_la_importacion_sin_guardar_nada(cliente_web, tmp_path: Path):
    raiz = tmp_path / "perfil"
    modulo_importacion.guardar_importacion(raiz, modulo_importacion.Importacion(skills=[_skill()]))
    cliente_web.post("/perfil/importar/descartar")
    assert modulo_importacion.cargar_importacion(raiz) is None
    assert almacen.cargar_perfil(raiz).skill("python") is None


def test_un_perfil_vacio_enlaza_a_importar_desde_mi_perfil(cliente_web):
    respuesta = cliente_web.get("/perfil")
    assert b'href="/perfil/importar"' in respuesta.data


def test_el_pie_de_cualquier_pantalla_enlaza_a_importar(cliente_web):
    respuesta = cliente_web.get("/ajustes")
    assert b'href="/perfil/importar"' in respuesta.data


def test_subir_un_formato_no_soportado_muestra_el_error(cliente_web):
    datos = {"fichero": (io.BytesIO(b"contenido cualquiera"), "cv.txt")}
    respuesta = cliente_web.post(
        "/perfil/importar", data=datos, content_type="multipart/form-data"
    )
    assert respuesta.status_code == 200
    assert "no es un formato soportado".encode("utf-8") in respuesta.data


# --------------------------------------------------------------------------
# De punta a punta: subir -> extraer -> analizar -> revisar -> guardar.
#
# Las pruebas de arriba verifican cada pieza por separado (la extracción, el
# análisis con un ClienteIA falso, el guardado de lo marcado); esta es la
# única que demuestra que la VISTA las conecta bien entre sí. Se sustituye
# `crear_cliente` por un doble (sin tocar la red ni una clave real) para
# poder probar la cadena completa sin gastar una llamada de verdad.
# --------------------------------------------------------------------------


class _ClienteFalsoDisponible:
    """Siempre "disponible" y siempre devuelve la misma propuesta — no hace
    falta variar la respuesta para demostrar que la vista sabe usarla."""

    def __init__(self, respuesta: str):
        self.respuesta = respuesta

    def completar(self, sistema: str, usuario: str) -> str:
        return self.respuesta

    def disponible(self) -> bool:
        return True


def _respuesta_ia() -> str:
    import json

    return json.dumps(
        {
            "experiencias": [
                {
                    "titulo": {"es": "Ingeniera de Datos", "en": "Data Engineer"},
                    "periodo": {"es": "2025 - actualidad", "en": "2025 - present"},
                    "bullets": {
                        "es": ["Pipeline de ingesta con Airflow"],
                        "en": ["Ingestion pipeline with Airflow"],
                    },
                    "stack": {"es": "Python, Airflow", "en": "Python, Airflow"},
                    "keywords": ["airflow", "etl"],
                }
            ],
            "skills": [
                {"nombre": {"es": "SQL", "en": "SQL"}, "categoria": "dato", "keywords": ["sql"]}
            ],
            "skills_personales": [
                {
                    "nombre": {"es": "Trabajo en equipo", "en": "Teamwork"},
                    "keywords": ["team player"],
                }
            ],
            "idiomas": [
                {
                    "nombre": {"es": "Inglés", "en": "English"},
                    "nivel": {"es": "C1 Avanzado", "en": "C1 Advanced"},
                    "keywords": ["advanced english"],
                }
            ],
        },
        ensure_ascii=False,
    )


def _docx_de_prueba() -> bytes:
    import docx

    documento = docx.Document()
    documento.add_paragraph("Ana Ejemplo — Ingeniera de Datos")
    documento.add_paragraph("2025 - actualidad: pipeline de ingesta con Airflow y Python.")
    buffer = io.BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def test_de_punta_a_punta_subir_analizar_revisar_y_guardar(cliente_web, tmp_path, monkeypatch):
    import cv_adaptativo.web.vistas.importar as vista_importar

    monkeypatch.setattr(
        vista_importar,
        "crear_cliente",
        lambda proveedor, clave: _ClienteFalsoDisponible(_respuesta_ia()),
    )

    # 1) Subir un .docx real.
    subida = cliente_web.post(
        "/perfil/importar",
        data={"fichero": (io.BytesIO(_docx_de_prueba()), "mi-cv.docx")},
        content_type="multipart/form-data",
    )
    assert subida.status_code == 302
    assert subida.location.endswith("/perfil/importar/revisar")

    # 2) La pantalla de revisión muestra lo que devolvió el análisis, no un
    #    placeholder ni datos de otra prueba.
    revision = cliente_web.get("/perfil/importar/revisar")
    html = revision.data.decode("utf-8")
    assert "Ingeniera de Datos" in html
    assert "Airflow" in html
    assert "SQL" in html
    assert "Trabajo en equipo" in html
    assert "Inglés" in html

    # 3) Guardar la experiencia y la skill personal; descartar la skill
    #    técnica y el idioma, para probar los dos caminos (guardar/descartar)
    #    en las cuatro categorías a la vez.
    guardar = cliente_web.post(
        "/perfil/importar/guardar",
        data={
            "exp-0": "1",
            "exp-0-titulo_es": "Ingeniera de Datos",
            "exp-0-titulo_en": "Data Engineer",
            "exp-0-periodo_es": "2025 - actualidad",
            "exp-0-periodo_en": "2025 - present",
            "exp-0-bullets_es": "Pipeline de ingesta con Airflow",
            "exp-0-bullets_en": "Ingestion pipeline with Airflow",
            "exp-0-stack_es": "Python, Airflow",
            "exp-0-stack_en": "Python, Airflow",
            # "skill-0" no se envía: queda descartada.
            "skillpersonal-0": "1",
            "skillpersonal-0-nombre_es": "Trabajo en equipo",
            "skillpersonal-0-nombre_en": "Teamwork",
            # "idioma-0" no se envía: queda descartado.
        },
    )
    assert guardar.status_code == 302

    # 4) El perfil real tiene la experiencia y la skill personal, y NO tiene
    #    ni la skill técnica ni el idioma descartados.
    perfil = almacen.cargar_perfil(tmp_path / "perfil")
    experiencias = [e for e in perfil.experiencias if e.titulo["es"] == "Ingeniera de Datos"]
    assert len(experiencias) == 1
    assert experiencias[0].bullets["es"] == ["Pipeline de ingesta con Airflow"]
    assert perfil.skill("sql") is None
    assert perfil.skill_personal("trabajo-en-equipo") is not None
    assert perfil.skill("trabajo-en-equipo") is None  # nunca en el catálogo técnico
    assert perfil.idioma("ingles") is None

    # 5) El borrador de importación queda limpio: no se puede volver a
    #    "revisar" la misma tanda ni dejar basura en el perfil.
    assert modulo_importacion.cargar_importacion(tmp_path / "perfil") is None
