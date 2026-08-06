"""Tests de la extracción determinista de texto (PDF y .docx).

Usa ficheros reales generados en el propio test —un PDF válido byte a byte,
un .docx real vía python-docx— en vez de mockear las librerías: lo que
importa aquí es que la extracción sea *fiel*, y eso solo lo demuestra pasar
por el parser de verdad.
"""
from __future__ import annotations

import io
import sys

import pytest

from ancla.perfil.extraccion import ErrorExtraccion, extraer_texto


def _pdf_minimo(texto: str) -> bytes:
    """Un PDF válido de una página con `texto` como contenido extraíble,
    construido con offsets de xref exactos (no aproximados)."""
    contenido = f"BT /F1 12 Tf 10 50 Td ({texto}) Tj ET".encode("latin-1")
    objetos = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>"
        b"/MediaBox[0 0 300 100]/Contents 5 0 R>>",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
        b"<</Length " + str(len(contenido)).encode() + b">>\nstream\n"
        + contenido + b"\nendstream",
    ]
    partes = [b"%PDF-1.4\n"]
    offsets = [0]
    for n, cuerpo in enumerate(objetos, start=1):
        offsets.append(sum(len(p) for p in partes))
        partes.append(f"{n} 0 obj".encode() + cuerpo + b"endobj\n")
    inicio_xref = sum(len(p) for p in partes)
    xref = [b"xref\n", f"0 {len(objetos) + 1}\n".encode(), b"0000000000 65535 f \n"]
    for offset in offsets[1:]:
        xref.append(f"{offset:010d} 00000 n \n".encode())
    partes += xref
    partes.append(
        f"trailer<</Size {len(objetos) + 1}/Root 1 0 R>>\n"
        f"startxref\n{inicio_xref}\n%%EOF".encode()
    )
    return b"".join(partes)


def _docx_con(parrafos: list[str]) -> bytes:
    import docx

    documento = docx.Document()
    for parrafo in parrafos:
        documento.add_paragraph(parrafo)
    buffer = io.BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def test_extrae_texto_de_un_pdf_real():
    datos = _pdf_minimo("Experiencia en Python y SQL, tres anhos en el sector.")
    texto = extraer_texto("cv.pdf", datos)
    assert "Experiencia en Python y SQL" in texto


def test_extrae_texto_de_un_docx_real():
    datos = _docx_con(
        [
            "Daniel Vega Rosado",
            "ML Developer — Telco Churn (2026 - actualidad)",
            "Pipeline completo de datos con Python y Optuna.",
        ]
    )
    texto = extraer_texto("cv.docx", datos)
    assert "ML Developer" in texto
    assert "Python y Optuna" in texto


def test_extrae_texto_de_tablas_en_docx():
    """Un CV suele meter el stack o las fechas en una tabla, no solo en párrafos."""
    import docx

    documento = docx.Document()
    documento.add_paragraph("Daniel Vega Rosado — Ingeniero de datos junior")
    tabla = documento.add_table(rows=1, cols=2)
    tabla.rows[0].cells[0].text = "Lenguajes: Python, SQL"
    tabla.rows[0].cells[1].text = "Herramientas: Optuna, Airflow"
    buffer = io.BytesIO()
    documento.save(buffer)

    texto = extraer_texto("cv.docx", buffer.getvalue())
    assert "Python, SQL" in texto and "Optuna, Airflow" in texto


def test_un_formato_no_soportado_lanza_error_claro():
    with pytest.raises(ErrorExtraccion, match="cv.txt"):
        extraer_texto("cv.txt", b"cualquier cosa")


def test_sin_pypdf_instalado_lo_dice_en_vez_de_culpar_al_pdf(monkeypatch):
    """El caso real que motivó esto: Daniel tenía `pypdf` instalado en un
    entorno distinto del que ejecuta `run.py`. `ModuleNotFoundError` es una
    `Exception` como cualquier otra, así que la red de seguridad genérica
    la atrapaba — pero con el mensaje de "tu PDF puede estar dañado", que es
    activamente engañoso cuando el problema real es una dependencia que
    falta. Tiene que salir el mismo mensaje accionable que ya usa
    `ia/groq.py` para el mismo caso con el SDK de Groq."""
    monkeypatch.setitem(sys.modules, "pypdf", None)

    with pytest.raises(ErrorExtraccion, match="pip install -r requirements.txt"):
        extraer_texto("cv.pdf", b"lo que sea")


def test_sin_python_docx_instalado_lo_dice_tambien(monkeypatch):
    monkeypatch.setitem(sys.modules, "docx", None)

    with pytest.raises(ErrorExtraccion, match="pip install -r requirements.txt"):
        extraer_texto("cv.docx", b"lo que sea")


def test_una_excepcion_no_prevista_del_extractor_no_llega_a_500(monkeypatch):
    """El caso real que motivó esto: un PDF real (protegido, con fuentes
    raras, exportado por una herramienta concreta) puede hacer fallar pypdf
    con una excepción distinta de `PdfReadError` — la única que se atrapaba
    dentro de `_texto_desde_pdf`. Como la vista solo captura `ErrorExtraccion`,
    cualquier otra cosa subía sin control hasta Flask y salía como 500.
    Aquí se simula ese "cualquier otra cosa" con un `RuntimeError` genérico,
    para no depender de un PDF patológico concreto y frágil."""
    import pypdf

    class LectorRoto:
        def __init__(self, *a, **kw):
            raise RuntimeError("fallo interno de pypdf no relacionado con el formato")

    monkeypatch.setattr(pypdf, "PdfReader", LectorRoto)

    with pytest.raises(ErrorExtraccion, match="No se ha podido leer"):
        extraer_texto("cv.pdf", b"lo que sea, no llega a importar")


def test_un_pdf_invalido_lanza_error_en_vez_de_reventar():
    with pytest.raises(ErrorExtraccion):
        extraer_texto("cv.pdf", b"esto no es un PDF de verdad")


def test_un_docx_invalido_lanza_error_en_vez_de_reventar():
    with pytest.raises(ErrorExtraccion):
        extraer_texto("cv.docx", b"esto no es un docx de verdad")


def test_un_pdf_sin_texto_real_avisa_de_que_puede_ser_una_imagen():
    """Un PDF escaneado (imagen sin capa de texto) extrae vacío o casi vacío:
    hay que decirlo, no fingir que se leyó un CV en blanco."""
    datos = _pdf_minimo("")
    with pytest.raises(ErrorExtraccion, match="escaneado"):
        extraer_texto("cv.pdf", datos)


def test_la_extension_no_distingue_mayusculas():
    datos = _pdf_minimo("Contenido de prueba suficientemente largo para pasar el minimo.")
    assert "Contenido de prueba" in extraer_texto("CV.PDF", datos)
