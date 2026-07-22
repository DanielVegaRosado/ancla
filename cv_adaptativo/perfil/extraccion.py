"""Extracción determinista de texto desde un CV subido (PDF o Word).

Nada de IA aquí a propósito: sacar el texto de un fichero es un problema
resuelto por una librería, no algo que necesite criterio. Que esta capa sea
100% mecánica es lo que hace fiable el paso siguiente (`importador.py`): el
modelo analiza texto garantizado idéntico al del fichero original, nunca una
lectura visual suya que podría transcribir mal una palabra o saltarse una
línea sin que nadie lo note.

Añadir un formato nuevo es una entrada más en `_EXTRACTORES`, no tocar la
función que despacha — abierto a extensión, cerrado a modificación.
"""
from __future__ import annotations

import io
from typing import Callable

MIN_CARACTERES_UTILES = 40


class ErrorExtraccion(Exception):
    """El fichero no se pudo leer, o no tenía texto aprovechable dentro.

    El mensaje va en español y pensado para enseñarse tal cual: es el mismo
    criterio que `ErrorPerfil` en el resto del perfil.
    """


def _texto_desde_pdf(datos: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        lector = PdfReader(io.BytesIO(datos))
        return "\n".join(pagina.extract_text() or "" for pagina in lector.pages)
    except PdfReadError as exc:
        raise ErrorExtraccion(
            "No se ha podido leer el PDF: parece estar dañado o protegido."
        ) from exc


def _texto_desde_docx(datos: bytes) -> str:
    import docx

    try:
        documento = docx.Document(io.BytesIO(datos))
    except Exception as exc:  # python-docx no expone una excepción propia
        raise ErrorExtraccion(
            "No se ha podido leer el documento: ¿es un .docx válido?"
        ) from exc
    partes = [parrafo.text for parrafo in documento.paragraphs]
    for tabla in documento.tables:
        for fila in tabla.rows:
            partes.extend(celda.text for celda in fila.cells)
    return "\n".join(partes)


_EXTRACTORES: dict[str, Callable[[bytes], str]] = {
    ".pdf": _texto_desde_pdf,
    ".docx": _texto_desde_docx,
}

FORMATOS_SOPORTADOS = tuple(_EXTRACTORES)


def extraer_texto(nombre_fichero: str, datos: bytes) -> str:
    """Extensión del fichero -> extractor. `ErrorExtraccion` si no se
    reconoce el formato, si no se puede leer, o si el resultado es
    demasiado corto para ser un CV (p. ej. un PDF escaneado sin texto real).

    Un PDF o .docx real de verdad —protegido, con fuentes raras, exportado
    por una herramienta que produce algo ligeramente distinto de lo que
    espera la librería— puede hacer fallar `pypdf`/`python-docx` con
    excepciones que no son las que ya se atrapan dentro de cada extractor
    concreto. La red de seguridad va aquí, en el punto de despacho, para
    cubrir cualquier extractor —también los que se añadan más adelante—:
    esta función nunca debe dejar pasar nada que no sea `ErrorExtraccion`.
    """
    extension = _extension(nombre_fichero)
    extractor = _EXTRACTORES.get(extension)
    if extractor is None:
        formatos = ", ".join(FORMATOS_SOPORTADOS)
        raise ErrorExtraccion(
            f"«{nombre_fichero}» no es un formato soportado. Admitidos: {formatos}."
        )

    try:
        texto = extractor(datos).strip()
    except ErrorExtraccion:
        raise  # Ya viene con un mensaje pensado para enseñarse tal cual.
    except Exception as exc:
        raise ErrorExtraccion(
            f"No se ha podido leer «{nombre_fichero}». Puede estar dañado, protegido con "
            "contraseña, o en una variante del formato que esta versión no reconoce. "
            "Si puedes, pega el texto a mano en su lugar."
        ) from exc

    if len(texto) < MIN_CARACTERES_UTILES:
        raise ErrorExtraccion(
            f"No se ha podido sacar texto de «{nombre_fichero}». Si es un PDF "
            "escaneado (una imagen, no texto real), pega el contenido a mano."
        )
    return texto


def _extension(nombre_fichero: str) -> str:
    punto = nombre_fichero.rfind(".")
    return nombre_fichero[punto:].lower() if punto != -1 else ""
