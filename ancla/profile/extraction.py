"""Deterministic text extraction from an uploaded CV (PDF or Word).

No AI here, on purpose: getting text out of a file is a problem already
solved by a library, not something that needs judgment. This layer being
100% mechanical is what makes the next step (`importador.py`) trustworthy:
the model analyses text guaranteed to be identical to the original file,
never its own visual reading, which could mistranscribe a word or skip a
line without anyone noticing.

Adding a new format is one more entry in `_EXTRACTORES`, never touching the
dispatch function — open for extension, closed for modification.
"""
from __future__ import annotations

import io
import zipfile
from typing import Callable

from flask_babel import gettext as _

from ancla.profile import zip_safety

MIN_CARACTERES_UTILES = 40
# A .docx is a zip: a crafted small file could claim a huge uncompressed
# size (zip bomb). MAX_CONTENT_LENGTH in web/__init__.py already bounds the
# compressed upload, but not what it expands to — this bounds that.
MAX_BYTES_DESCOMPRIMIDOS_DOCX = 100 * 1024 * 1024  # 100 MB


class ExtractionError(Exception):
    """The file could not be read, or had no usable text inside.

    The message is in Spanish and meant to be shown as-is: same approach as
    `ErrorPerfil` across the rest of the profile.
    """


def _text_from_pdf(datos: bytes) -> str:
    # Lazy import, same as `ia/groq.py`: this way the app still starts and
    # can show the message below even if the dependency is not installed,
    # instead of dying at startup on a module-level import.
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError as exc:
        raise ExtractionError(
            _("Falta la librería «pypdf». Instálala con: pip install -r requirements.txt")
        ) from exc

    try:
        lector = PdfReader(io.BytesIO(datos))
        return "\n".join(pagina.extract_text() or "" for pagina in lector.pages)
    except PdfReadError as exc:
        raise ExtractionError(
            _("No se ha podido leer el PDF: parece estar dañado o protegido.")
        ) from exc


def _text_from_docx(datos: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ExtractionError(
            _("Falta la librería «python-docx». Instálala con: pip install -r requirements.txt")
        ) from exc

    try:
        with zipfile.ZipFile(io.BytesIO(datos)) as zip_:
            tamano_descomprimido = zip_safety.uncompressed_size(zip_)
    except zipfile.BadZipFile as exc:
        raise ExtractionError(
            _("No se ha podido leer el documento: ¿es un .docx válido?")
        ) from exc
    if tamano_descomprimido > MAX_BYTES_DESCOMPRIMIDOS_DOCX:
        raise ExtractionError(
            _("El documento «.docx» es sospechosamente grande al descomprimirlo. No se ha procesado.")
        )

    try:
        documento = docx.Document(io.BytesIO(datos))
    except Exception as exc:  # python-docx does not expose an exception of its own
        raise ExtractionError(
            _("No se ha podido leer el documento: ¿es un .docx válido?")
        ) from exc
    partes = [parrafo.text for parrafo in documento.paragraphs]
    for tabla in documento.tables:
        for fila in tabla.rows:
            partes.extend(celda.text for celda in fila.cells)
    return "\n".join(partes)


_EXTRACTORES: dict[str, Callable[[bytes], str]] = {
    ".pdf": _text_from_pdf,
    ".docx": _text_from_docx,
}

FORMATOS_SOPORTADOS = tuple(_EXTRACTORES)


def extract_text(nombre_fichero: str, datos: bytes) -> str:
    """File extension -> extractor. Raises `ErrorExtraccion` if the format is
    not recognised, if it cannot be read, or if the result is too short to be
    a CV (e.g. a scanned PDF with no real text).

    A genuinely real-world PDF or .docx — protected, with unusual fonts,
    exported by a tool that produces something slightly different from what
    the library expects — can make `pypdf`/`python-docx` fail with
    exceptions that are not the ones already caught inside each specific
    extractor. The safety net lives here, at the dispatch point, to cover
    any extractor — including ones added later: this function must never let
    anything through that is not `ErrorExtraccion`.
    """
    extension = _extension(nombre_fichero)
    extractor = _EXTRACTORES.get(extension)
    if extractor is None:
        formatos = ", ".join(FORMATOS_SOPORTADOS)
        raise ExtractionError(
            _(
                "«%(nombre)s» no es un formato soportado. Admitidos: %(formatos)s.",
                nombre=nombre_fichero,
                formatos=formatos,
            )
        )

    try:
        texto = extractor(datos).strip()
    except ExtractionError:
        raise  # Already comes with a message meant to be shown as-is.
    except Exception as exc:
        raise ExtractionError(
            _(
                "No se ha podido leer «%(nombre)s». Puede estar dañado, protegido con "
                "contraseña, o en una variante del formato que esta versión no reconoce. "
                "Si puedes, pega el texto a mano en su lugar.",
                nombre=nombre_fichero,
            )
        ) from exc

    if len(texto) < MIN_CARACTERES_UTILES:
        raise ExtractionError(
            _(
                "No se ha podido sacar texto de «%(nombre)s». Si es un PDF "
                "escaneado (una imagen, no texto real), pega el contenido a mano.",
                nombre=nombre_fichero,
            )
        )
    return texto


def _extension(nombre_fichero: str) -> str:
    punto = nombre_fichero.rfind(".")
    return nombre_fichero[punto:].lower() if punto != -1 else ""
