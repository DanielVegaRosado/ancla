"""Fills a `.docx` template with a proposal's content.

Deterministic, no AI call: this is what guarantees rule 2 (never rewrite
the user) by architecture rather than by instruction, the same way
`perfil/extraccion.py` does for imports. `docxtpl` only substitutes the
tags already in the template's own runs, so its fonts, colors, columns and
tables survive untouched.

`Experience` has no separate company field (see `ancla/profile/model.py`):
`title` already bundles "role · company" together, the same way it is
shown everywhere else in the app (Proposal screen, `proposal/format.py`).
`empresa` is kept in the context because the field catalog names it, but is
always empty — a template that wants role and company on separate lines has
nothing to split them from and has to fold both into `puesto`.

Each experience also carries `stack` (the tech-stack line, e.g. "Python,
FastAPI, PostgreSQL"): not in the original field catalog, added because a
real template (`corporativa-clasica.docx`) shows it under every experience
and the data was already sitting unused on `Experience.stack`. Optional —
a template that ignores the tag simply doesn't show it.

`contacto`, `titular`, `educacion` and `foto` mirror the "always complete,
never selected" rule that already governs `skills_personales`/`idiomas`:
none of the four go through `selection/engine.py`, so they are read live
from the profile, not from `Proposal`.
"""
from __future__ import annotations

import math
from io import BytesIO
from pathlib import Path

from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docxtpl import DocxTemplate, InlineImage

from ancla.export import metricas_aileron, metricas_montserrat
from ancla.export.templates import ExportTemplate, TemplateGeometry
from ancla.profile import store
from ancla.profile.model import Experience, Language, Profile, Proposal, SelectedExperience
from ancla.proposal.format import language_lines, personal_skill_names, skill_names

# Width and height of the embedded photo. Equal on purpose: `_recortar_a_circulo`
# below only turns a *square* bounding box into a visual circle. A portrait
# photo (the common case) gets its top/bottom cropped to a centered square
# before that — see that function's docstring.
LADO_FOTO = Mm(30)

# Section headers are plain text baked into each template (not tags a
# template author fills in), so they used to be permanently Spanish even
# when the proposal itself was generated in English. `propuesta.language`
# already tells us which; this is the ES/EN pair for each header a real
# template uses. Both real templates carry a slightly different label for
# the same concept ("SKILLS TÉCNICAS" vs "SKILLS", "EXPERIENCIA RELEVANTE"
# vs "EXPERIENCIA") — kept as separate keys so each template's own wording
# survives, rather than forcing both to share one generic label.
_ETIQUETAS: dict[str, dict[Language, str]] = {
    "contacto": {"es": "CONTACTO", "en": "CONTACT"},
    "educacion": {"es": "EDUCACIÓN", "en": "EDUCATION"},
    "skills_tecnicas": {"es": "SKILLS TÉCNICAS", "en": "TECHNICAL SKILLS"},
    "skills": {"es": "SKILLS", "en": "SKILLS"},
    "skills_personales": {"es": "SKILLS PERSONALES", "en": "PERSONAL SKILLS"},
    # Standalone subtitles ("TECHNICAL", "PERSONAL"), without repeating
    # "SKILLS" — for a template that groups one "SKILLS" header with both
    # of these underneath, instead of two independent headers.
    "tecnicas": {"es": "TÉCNICAS", "en": "TECHNICAL"},
    "personales": {"es": "PERSONALES", "en": "PERSONAL"},
    "idiomas": {"es": "IDIOMAS", "en": "LANGUAGES"},
    "sobre_mi": {"es": "SOBRE MÍ", "en": "ABOUT ME"},
    "experiencia_relevante": {"es": "EXPERIENCIA RELEVANTE", "en": "RELEVANT EXPERIENCE"},
    "experiencia": {"es": "EXPERIENCIA", "en": "EXPERIENCE"},
}


def _etiquetas(idioma: Language) -> dict[str, str]:
    return {f"etiqueta_{clave}": valores[idioma] for clave, valores in _ETIQUETAS.items()}


def resolved_selection(propuesta: Proposal, perfil: Profile) -> list[tuple[SelectedExperience, Experience]]:
    """The proposal's chosen experiences, in the same relevance order,
    paired with the profile data they point to. One that no longer exists
    there is dropped rather than shown as a blank block — same rule
    `proposal/format.py` follows for the copy-paste text."""
    parejas = ((seleccionada, perfil.experience(seleccionada.id)) for seleccionada in propuesta.experiences)
    return [(seleccionada, experiencia) for seleccionada, experiencia in parejas if experiencia is not None]


def resolved_experiences(propuesta: Proposal, perfil: Profile) -> list[Experience]:
    return [experiencia for _, experiencia in resolved_selection(propuesta, perfil)]


#: Non-breaking hyphen (U+2011), sometimes present in pasted text (e.g. a
#: "H‑Z‑H" bullet). Several export fonts ship without that glyph, so Word
#: silently falls back to a different font for just that character — it
#: reads as a font glitch mid-word. Swapped for a plain hyphen (U+002D,
#: visually identical) only in the exported `.docx`; the profile itself is
#: never touched, so this doesn't cross rule 2 ("never rewrite the user").
_GUION_NO_SEPARABLE = "‑"


def _sin_guion_no_separable(texto: str) -> str:
    return texto.replace(_GUION_NO_SEPARABLE, "-")




def build_context(
    propuesta: Proposal,
    perfil: Profile,
    experiencias: list[Experience],
    nombre: str,
    foto: InlineImage | str = "",
) -> dict:
    idioma = propuesta.language
    nombre_primero, _, nombre_resto = nombre.partition(" ")
    return {
        **_etiquetas(idioma),
        "nombre": nombre,
        # Some templates set the first name and the rest in different styles
        # ("**DANIEL** VEGA"), which a single tag cannot express: formatting
        # lives on the run, not on the text. Split on the first space; a
        # template that doesn't need it keeps using `nombre` whole.
        "nombre_primero": nombre_primero,
        "nombre_resto": nombre_resto,
        "sobre_mi": _sin_guion_no_separable(propuesta.about_me.text),
        "experiencias": [
            {
                "puesto": _sin_guion_no_separable(experiencia.title[idioma]),
                "empresa": "",
                "fechas": experiencia.period[idioma],
                "bullets": [_sin_guion_no_separable(b) for b in experiencia.bullets[idioma]],
                "stack": _sin_guion_no_separable(experiencia.stack[idioma]),
            }
            for experiencia in experiencias
        ],
        "skills": skill_names(propuesta, perfil),
        "idiomas": language_lines(perfil, idioma),
        "skills_personales": personal_skill_names(perfil, idioma),
        "contacto": [_sin_guion_no_separable(linea) for linea in perfil.contact],
        "titular": _sin_guion_no_separable(perfil.headline[idioma]),
        "educacion": [
            {
                "titulo": _sin_guion_no_separable(entrada.title[idioma]),
                "centro": _sin_guion_no_separable(entrada.institution[idioma]),
                "fechas": entrada.period[idioma],
            }
            for entrada in perfil.education
        ],
        "foto": foto,
    }


def render(
    plantilla: ExportTemplate,
    propuesta: Proposal,
    perfil: Profile,
    experiencias: list[Experience],
    nombre: str,
    root: Path,
) -> bytes:
    """Renders the template with this content and returns the finished
    `.docx`'s raw bytes, ready to stream as a download.

    `root` is only needed to resolve the profile photo's actual file
    (`Profile.photo` is just a filename — see `perfil/modelo.py`); an
    `InlineImage` has to be built against this specific `DocxTemplate`
    instance, which is why the photo can't be added in `build_context` on
    its own.
    """
    documento = DocxTemplate(plantilla.path)
    ruta_foto = store.photo_path(root)
    foto = InlineImage(documento, str(ruta_foto), width=LADO_FOTO, height=LADO_FOTO) if ruta_foto is not None else ""
    contexto = build_context(propuesta, perfil, experiencias, nombre, foto)
    documento.render(contexto)
    if ruta_foto is not None:
        _recortar_a_circulo(documento, ruta_foto)
    if plantilla.geometry is not None:
        _ajustar_densidad_experiencias(documento, contexto, plantilla.geometry)
    buffer = BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


# A template's experience block is laid out tight enough on paper that its
# own declared `capacidad_experiencias` (`ExportTemplate.capacity_experiences`
# — the most a design was built for) fits on one page — see
# `docx-templates/README.md`. With fewer experiences than that, a fixed
# density left the page looking short: a factor based only on how many
# experiences are missing cannot tell a short bullet from a long one, so it
# either overshoots or undershoots depending on the actual text. Instead,
# this estimates from the real bullet text how many lines the block will
# wrap to, and solves for the exact spacing/font-size factor that fills the
# page. The actual page measurements it needs to do that (column width,
# base font sizes, margins...) come from `ExportTemplate.geometry`, which
# each template's own build script writes to its `.yaml` — never
# hardcoded here, so there is only one place per template where those
# numbers are typed by hand.

# Loosening spacing alone has a ceiling: with genuinely sparse content (a
# handful of experiences with short bullets each), closing the gap to the
# page margin through blank space alone needs an unnaturally large
# line-spacing multiplier, which reads as visibly broken. So the scale
# factor also grows the font size a little, the way a person adjusting a
# CV by hand would: more text per line, and more lines that wrap, not just
# more air.
_ESCALA_MAXIMA = 2.2
# The other direction of the same idea: real bullets and a real "About me"
# can add up to more than the page has room for even at the template's own
# capacity (a long bullet, a long bio) — tightening spacing and font size a
# little is what a person would do by hand rather than letting the CV spill
# onto a second page for two lines. `escala` is never pushed below this so
# text never becomes hard to read chasing an exact fit.
_ESCALA_MINIMA = 0.75
_SENSIBILIDAD_TAMANO = 0.35  # how much the font grows or shrinks per point of `escala`


# One character-width table set per font a template might use, keyed by
# the exact name `ExportTemplate.geometry.fuente` declares — different
# fonts have different average glyph widths (Montserrat is noticeably
# wider than Aileron), so the line-wrap estimate below has to use the
# right one, not whichever happened to be imported first. A font missing
# here falls back to Aileron's table instead of crashing: a slightly-off
# density scale is a cosmetic risk, not a data-safety one.
_METRICAS_POR_FUENTE: dict[str, tuple[dict[str, float], dict[str, float], dict[str, float]]] = {
    "Aileron": (metricas_aileron.ANCHOS_REGULAR, metricas_aileron.ANCHOS_BOLD, metricas_aileron.ANCHOS_ITALIC),
    "Montserrat": (
        metricas_montserrat.ANCHOS_REGULAR, metricas_montserrat.ANCHOS_BOLD, metricas_montserrat.ANCHOS_ITALIC
    ),
}


def _metricas(geometry: TemplateGeometry) -> tuple[dict[str, float], dict[str, float], dict[str, float]]:
    return _METRICAS_POR_FUENTE.get(geometry.fuente, _METRICAS_POR_FUENTE["Aileron"])


def _ancho_em(texto: str, tabla: dict[str, float]) -> float:
    return sum(tabla.get(caracter, 0.5) for caracter in texto)


def _lineas_estimadas(texto: str, tabla: dict[str, float], tamano_pt: float, ancho_columna: float) -> int:
    if not texto:
        return 0
    ancho_pt = _ancho_em(texto, tabla) * tamano_pt
    return max(1, math.ceil(ancho_pt / ancho_columna))


def _factor_tamano(escala: float) -> float:
    return 1 + (escala - 1) * _SENSIBILIDAD_TAMANO


def _altura_bloque(experiencia: dict, escala: float, geometry: TemplateGeometry) -> float:
    factor_tamano = _factor_tamano(escala)
    s_titulo = geometry.tamano_titulo * factor_tamano
    s_cuerpo = geometry.tamano_cuerpo * factor_tamano
    s_stack = geometry.tamano_stack * factor_tamano
    multiplicador_linea = 1 + (geometry.linea_bullet - 1) * min(escala, 1.3)
    anchos_regular, anchos_bold, anchos_italic = _metricas(geometry)

    lineas_titulo = _lineas_estimadas(experiencia["puesto"], anchos_bold, s_titulo, geometry.ancho_columna_titulo)
    altura = lineas_titulo * s_titulo + geometry.espacio_titulo * escala

    for bullet in experiencia["bullets"]:
        lineas = _lineas_estimadas(bullet, anchos_regular, s_cuerpo, geometry.ancho_columna)
        altura += lineas * s_cuerpo * multiplicador_linea + geometry.espacio_bullet * escala

    if experiencia["stack"]:
        lineas = _lineas_estimadas(experiencia["stack"], anchos_italic, s_stack, geometry.ancho_columna)
        altura += lineas * s_stack * multiplicador_linea + geometry.espacio_bullet * escala

    return altura * geometry.calibracion


def _altura_texto(texto: str, tamano: float, linea: float, geometry: TemplateGeometry, *, negrita: bool = False) -> float:
    """Real height of a paragraph of never-scaled text, estimated from its
    actual content the same way an experience bullet is — instead of
    assuming a fixed average length for it. Used for anything whose length
    genuinely varies per profile or proposal (a name, a role, an "About
    me" paragraph...), never for a template's own fixed section-header
    labels, which are always short and never wrap. `tamano <= 0` means this
    template has no such text in this particular slot."""
    if not texto or tamano <= 0:
        return 0.0
    anchos_regular, anchos_bold, _ = _metricas(geometry)
    lineas = _lineas_estimadas(texto, anchos_bold if negrita else anchos_regular, tamano, geometry.ancho_columna)
    return lineas * tamano * linea * geometry.calibracion


def _altura_cabecera(contexto: dict, geometry: TemplateGeometry) -> float:
    """Real height of the name + role block above the experience section —
    both vary per profile (a long name or role can wrap), so neither is
    assumed to take a fixed amount of room.

    Some designs (Minimalista Cálida) stack the first and last name on two
    differently-sized lines; others (Corporativa Clásica) put both on one
    shared line at one size. `tamano_nombre_resto <= 0` is how a template
    says "no second line" — same 0-disables-this-slot convention as the
    rest of `TemplateGeometry` — so the two names are measured combined,
    as the one line they actually share, instead of double-counted as two.
    """
    primero, resto = contexto["nombre_primero"], contexto["nombre_resto"]
    if geometry.tamano_nombre_resto > 0:
        altura_nombre = _altura_texto(
            primero, geometry.tamano_nombre_primero, 1.0, geometry, negrita=True
        ) + _altura_texto(resto, geometry.tamano_nombre_resto, 1.0, geometry, negrita=True)
    else:
        altura_nombre = _altura_texto(
            f"{primero} {resto}".strip(), geometry.tamano_nombre_primero, 1.0, geometry, negrita=True
        )
    return altura_nombre + _altura_texto(contexto["titular"], geometry.tamano_titular, 1.0, geometry)


def _escala_necesaria(experiencias: list[dict], geometry: TemplateGeometry, contexto: dict) -> float:
    """The scale (spacing + font size) that makes the experience block
    reach exactly the page's bottom margin, calculated from the real text —
    not from how many experiences there are, which alone says nothing about
    how much space they take up (a long bullet weighs more than a short
    one). Every other real-content block sharing the page (name, role,
    "About me", wherever it sits) is measured the same way, so none of them
    is guessed either.

    Can come back below 1.0: real content can add up to more than the page
    has room for even at a template's own capacity, and shrinking a little
    is the fix, the same way growing is the fix for too little — this
    solves for the single `escala` that reaches the target from either
    side, not just the "not enough content" one."""
    if not experiencias:
        return 1.0
    sobre_mi = contexto.get("sobre_mi", "")
    objetivo = (
        geometry.alto_pagina
        - geometry.margen_inferior_seguridad
        - geometry.altura_hasta_experiencia
        - _altura_cabecera(contexto, geometry)
        - _altura_texto(sobre_mi, geometry.tamano_texto_antes_experiencia, geometry.linea_texto_antes_experiencia, geometry)
        - geometry.altura_tras_experiencia
        - _altura_texto(sobre_mi, geometry.tamano_texto_tras_experiencia, geometry.linea_texto_tras_experiencia, geometry)
    )

    def altura_total(escala: float) -> float:
        return sum(_altura_bloque(e, escala, geometry) for e in experiencias)

    # `altura_total` grows with `escala` in both directions (more/less
    # space, more/less letter size), so a binary search finds the exact
    # scale without inverting the formula by hand — valid even if the
    # height calculation changes.
    bajo, alto = _ESCALA_MINIMA, _ESCALA_MAXIMA
    for _ in range(20):
        medio = (bajo + alto) / 2
        if altura_total(medio) < objetivo:
            bajo = medio
        else:
            alto = medio
    return alto


def _aflojar_espaciado(paragraph, escala: float, *, escalar_tamano: bool = True, tamano_base: float | None = None) -> None:
    """Scales this paragraph's spacing and, unless `escalar_tamano` is
    False, its runs' font size. The exemption exists for text in a
    fixed-width column that must never wrap to a second line — growing its
    font with the rest of the page would defeat that constraint.

    `run.font.size` reads back as `None` — no explicit size, not a zero —
    whenever a run's size happens to match its paragraph/style's inherited
    default: Word treats the explicit value as redundant and strips it the
    next time the file is opened and saved (e.g. during font embedding in
    `herramientas/incrustar_fuentes.py`), even though the rendered text
    looks identical either way. That silently broke scaling here — the run
    was real and visible, but looked unset, so it was skipped. `tamano_base`
    is what the caller already knows this run's real size should be from
    the template's own geometry, used only to fill that gap; a run that
    still carries its own explicit size keeps using it, never `tamano_base`.
    """
    factor_tamano = _factor_tamano(escala)
    pf = paragraph.paragraph_format
    if pf.space_before is not None:
        pf.space_before = Pt(pf.space_before.pt * escala)
    if pf.space_after is not None:
        pf.space_after = Pt(pf.space_after.pt * escala)
    if isinstance(pf.line_spacing, (int, float)):
        # Line spacing grows more slowly than paragraph spacing: a high
        # line-spacing multiplier stands out much faster than a few extra
        # points before/after a paragraph.
        pf.line_spacing = 1 + (pf.line_spacing - 1) * min(escala, 1.3)
    if not escalar_tamano:
        return
    for run in paragraph.runs:
        tamano_actual = run.font.size.pt if run.font.size is not None else tamano_base
        if tamano_actual is not None:
            run.font.size = Pt(tamano_actual * factor_tamano)


def _ajustar_densidad_experiencias(documento: DocxTemplate, contexto: dict, geometry: TemplateGeometry) -> None:
    """Loosens the experience block to `_escala_necesaria`, the exact
    factor that makes it reach the page's bottom margin. A no-op for any
    template without the expected table structure (a main column holding a
    nested, one-row-per-experience table) — a template opts into this
    adjustment simply by declaring `geometria` in its `.yaml`, and this
    still degrades gracefully if its actual layout doesn't match."""
    experiencias = contexto["experiencias"]
    if not experiencias:
        return
    escala = _escala_necesaria(experiencias, geometry, contexto)
    if escala == 1.0:
        return
    try:
        columna_principal = documento.docx.tables[0].rows[0].cells[1]
        tabla_experiencias = columna_principal.tables[0]
    except (IndexError, AttributeError):
        return
    for fila in tabla_experiencias.rows:
        celda = fila.cells[0]
        for tabla_cabecera in celda.tables:
            for fila_cabecera in tabla_cabecera.rows:
                # The date cell (second column: title | date) is exempt from
                # font-size scaling — its fixed width (`DATE_W`) is already
                # calibrated to the limit for a single line, and growing it
                # with density would make it wrap to two.
                for indice, celda_cabecera in enumerate(fila_cabecera.cells):
                    es_fecha = indice == 1
                    for parrafo in celda_cabecera.paragraphs:
                        _aflojar_espaciado(
                            parrafo, escala, escalar_tamano=not es_fecha, tamano_base=geometry.tamano_titulo
                        )
        for parrafo in celda.paragraphs:
            _aflojar_espaciado(parrafo, escala, tamano_base=geometry.tamano_cuerpo)


_NS_DIBUJO = "http://schemas.openxmlformats.org/drawingml/2006/main"

# `pic:cNvPr/@descr` value that marks a picture as decorative (e.g. a
# full-bleed background rectangle), not the profile photo. Templates that
# carry one set it so `_recortar_a_circulo` skips it — see the docstring.
MARCA_DECORATIVA = "ancla:decorativo"


def _recortar_a_circulo(documento: DocxTemplate, ruta_foto: Path) -> None:
    """Makes the just-embedded photo round, the way the two real templates'
    Canva originals show it.

    Two steps, because a non-square photo (almost every real headshot) would
    otherwise render as an oval: (1) crop to a centered square via the
    picture's `a:srcRect` — the fill's visible area, not the file on disk —
    computed from the source image's actual pixel dimensions; (2) swap the
    picture's shape geometry from `rect` to `ellipse`, which is what turns a
    square image into a circle on screen.

    A template may also carry other, purely decorative `pic:pic` elements
    (e.g. a full-bleed background rectangle anchored behind the text) —
    those are marked with `MARCA_DECORATIVA` in their `pic:cNvPr` `descr`
    so this loop can skip them instead of assuming every picture is the
    photo.
    """
    imagen = DocxImage.from_file(str(ruta_foto))
    ancho_px, alto_px = imagen.px_width, imagen.px_height
    # `documento.get_docx()` does NOT work here — it calls `init_docx()`,
    # which after a `render()` has already run reloads the original
    # template from disk and discards everything just rendered.
    # `documento.element` goes through `DocxTemplate.__getattr__`, which
    # delegates to `self.docx` without reloading.
    cuerpo = documento.element.body
    for figura in cuerpo.iter(qn("pic:pic")):
        nv_pr = figura.find(f".//{qn('pic:cNvPr')}")
        if nv_pr is not None and nv_pr.get("descr") == MARCA_DECORATIVA:
            continue
        geometria = figura.find(f".//{{{_NS_DIBUJO}}}prstGeom")
        if geometria is not None:
            geometria.set("prst", "ellipse")

        relleno = figura.find(qn("pic:blipFill"))
        if relleno is None:
            continue
        recorte = OxmlElement("a:srcRect")
        if ancho_px > alto_px:
            sobra = int((1 - alto_px / ancho_px) / 2 * 100000)
            recorte.set("l", str(sobra))
            recorte.set("r", str(sobra))
        elif alto_px > ancho_px:
            sobra = int((1 - ancho_px / alto_px) / 2 * 100000)
            recorte.set("t", str(sobra))
            recorte.set("b", str(sobra))
        else:
            continue  # already square, nothing to crop
        relleno.insert(1, recorte)  # right after <a:blip>, before <a:stretch>
